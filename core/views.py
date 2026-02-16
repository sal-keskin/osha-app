from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import HttpResponse, JsonResponse, FileResponse
from django.db.models import Q
from django.views.decorators.http import require_POST
import json
from .decorators import medical_access_required
import csv
import openpyxl
from datetime import datetime
from django.forms import inlineformset_factory
from .forms import (
    LoginForm, WorkplaceForm, WorkerForm, ProfessionalForm,
    EducationForm, InspectionForm, ExaminationForm, ProfessionForm,
    ExaminationNoteForm, FacilityForm, CustomUserCreationForm, CustomUserChangeForm,
    CertificateTemplateForm, AssessmentSessionForm, RiskToolForm, RiskToolImportForm, CustomRiskForm
)
from .models import (
    Workplace, Worker, Professional, Education, Inspection, Examination, Profession, Facility, ActionLog, CertificateTemplate,
    RiskTool, RiskCategory, RiskTopic, RiskQuestion, AssessmentSession, AssessmentCustomRisk, AssessmentAnswer, ActionPlanMeasure,
    RiskAssessmentTeamMember, RiskControlRecord, UserProfile, WorkplaceAssignment
)

from .utils import get_allowed_workplaces
from .stats import get_user_scoped_stats
from encrypted_model_fields.fields import EncryptedCharField, EncryptedTextField, EncryptedDateField, EncryptedBooleanField
from django.contrib.auth.models import User
# Removed duplicate imports
from .import_utils import ImportHandler
import json
from .pdf_generator import generate_certificate_pdf

def log_action(user, action, model_obj, details=None):
    if not user.is_authenticated: return
    try:
        model_name = model_obj._meta.verbose_name
        object_id = str(model_obj.pk)
        ActionLog.objects.create(
            user=user,
            action=action,
            model_name=model_name,
            object_id=object_id,
            details=details or str(model_obj)
        )
    except: pass

@login_required
def api_create_profession(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            name = data.get('name')
            if not name:
                return JsonResponse({'success': False, 'error': 'Name is required'})

            profession, created = Profession.objects.get_or_create(name=name)
            return JsonResponse({
                'success': True,
                'profession': {'id': profession.id, 'name': profession.name}
            })
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    return JsonResponse({'success': False, 'error': 'Invalid method'})

def login_view(request):
    if request.method == 'POST':
        form = LoginForm(request.POST, request=request)
        if form.is_valid():
            username = form.cleaned_data['username']
            password = form.cleaned_data['password']
            user = authenticate(request, username=username, password=password)
            if user is not None:
                login(request, user)
                ActionLog.objects.create(user=user, action='Giriş', model_name='Oturum', details='Kullanıcı giriş yaptı.')
                # Clear captcha session data
                if 'captcha_question' in request.session:
                    del request.session['captcha_question']
                if 'captcha_answer' in request.session:
                    del request.session['captcha_answer']
                return redirect('dashboard')
            else:
                messages.error(request, 'Hatalı kullanıcı adı veya şifre.')
    else:
        # Reset captcha on new load
        if 'captcha_answer' in request.session:
             del request.session['captcha_answer']
        form = LoginForm(request=request)
    
    return render(request, 'core/login.html', {'form': form})

def logout_view(request):
    if request.user.is_authenticated:
        ActionLog.objects.create(user=request.user, action='Çıkış', model_name='Oturum', details='Kullanıcı çıkış yaptı.')
    logout(request)
    return redirect('login')

@login_required
def get_workers_json(request):
    workplace_id = request.GET.get('workplace_id')
    workers = []
    if workplace_id:
        # Cannot use .values() with encrypted fields as it returns raw bytes
        # Must fetch objects to trigger partial decryption
        worker_objs = Worker.objects.filter(workplace_id=workplace_id).only('id', 'name', 'tckn')
        workers = [{'id': w.id, 'name': w.name, 'tckn': w.tckn} for w in worker_objs]
    return JsonResponse({'workers': workers})

@login_required
def api_get_facilities(request):
    workplace_id = request.GET.get('workplace_id')
    facilities = []
    if workplace_id:
        facilities = list(Facility.objects.filter(workplace_id=workplace_id).values('id', 'name'))
    return JsonResponse({'facilities': facilities})

@login_required
def api_search_nace(request):
    """Search NACE codes from JSON file"""
    import os
    from django.conf import settings
    
    query = request.GET.get('q', '').lower().strip()
    if len(query) < 2:
        return JsonResponse({'results': [], 'count': 0})
    
    # Load NACE codes from JSON file
    json_path = os.path.join(settings.BASE_DIR, 'nace-codes.json')
    
    try:
        with open(json_path, 'r', encoding='utf-8') as f:
            nace_codes = json.load(f)
    except FileNotFoundError:
        return JsonResponse({'results': [], 'count': 0, 'error': 'NACE codes file not found'})
    
    # Search in both nace_code and description
    results = []
    for item in nace_codes:
        if (query in item.get('nace_code', '').lower() or 
            query in item.get('description', '').lower()):
            results.append({
                'nace_code': item.get('nace_code'),
                'description': item.get('description'),
                'danger_class': item.get('danger_class')
            })
            if len(results) >= 10:  # Limit results
                break
    
    return JsonResponse({'results': results, 'count': len(results)})


@login_required
def dashboard(request):
    from datetime import date, timedelta
    
    # 1. Apply Scope
    allowed_workplaces = get_allowed_workplaces(request.user)
    
    # 2. Filter Data based on Scope
    workers = Worker.objects.filter(workplace__in=allowed_workplaces)\
                            .select_related('workplace')\
                            .prefetch_related('education_set', 'examination_set')

    # Attention Logic
    missing_edu = 0
    missing_exam = 0
    expiring_edu = 0
    expiring_exam = 0
    today = date.today()

    for w in workers:
        # Education
        edus = w.education_set.all()
        if not edus:
            missing_edu += 1
        else:
            latest_edu = max(edus, key=lambda x: x.date, default=None)
            if latest_edu:
                validity = w.workplace.get_validity_years('education')
                expiry = latest_edu.date + timedelta(days=365*validity)
                if expiry.year == today.year and expiry.month == today.month:
                     expiring_edu += 1

        # Examination
        exams = w.examination_set.all()
        if not exams:
            missing_exam += 1
        else:
            latest_exam = max(exams, key=lambda x: x.date, default=None)
            if latest_exam:
                validity = w.workplace.get_validity_years('examination')
                expiry = latest_exam.date + timedelta(days=365*validity)
                if expiry.year == today.year and expiry.month == today.month:
                     expiring_exam += 1

    attention_data = {
        'missing_edu': missing_edu,
        'missing_exam': missing_exam,
        'expiring_edu': expiring_edu,
        'expiring_exam': expiring_exam,
        'has_issues': any([missing_edu, missing_exam, expiring_edu, expiring_exam])
    }

    # Context with Scoped Counts
    context = {
        'workplace_count': allowed_workplaces.count(),
        'facility_count': Facility.objects.filter(workplace__in=allowed_workplaces).count(),
        'worker_count': workers.count(),
        'professional_count': Professional.objects.count(), # Professionals are global
        'education_count': Education.objects.filter(workplace__in=allowed_workplaces).count(),
        'inspection_count': Inspection.objects.filter(workplace__in=allowed_workplaces).count(),
        'examination_count': Examination.objects.filter(worker__workplace__in=allowed_workplaces).count(),
        'risk_tool_count': RiskTool.objects.count(), # Global Templates
        'assessment_count': AssessmentSession.objects.filter(facility__workplace__in=allowed_workplaces).count(),
        'attention': attention_data,
    }
    return render(request, 'core/dashboard.html', context)

# Filtering Helper
# Filtering Helper
def apply_filters(queryset, filter_config, params):
    """
    Applies filters to the queryset based on configuration and request parameters.
    Handles encrypted fields by switching to Python-side filtering.
    """
    is_list = isinstance(queryset, list)
    model = queryset.model if hasattr(queryset, 'model') else (queryset[0].__class__ if queryset else None)

    for config in filter_config:
        field_name = config['field']
        param_value = params.get(field_name)
        
        # Check if field is encrypted
        is_encrypted = False
        if model:
            try:
                field = model._meta.get_field(field_name)
                if isinstance(field, (EncryptedCharField, EncryptedTextField, EncryptedDateField, EncryptedBooleanField)):
                    is_encrypted = True
            except: pass

        if param_value:
            # Update config with the current value
            config['value'] = param_value
            filter_type = config.get('type', 'text')

            if is_encrypted or is_list:
                # Python filtering
                if not is_list:
                    queryset = list(queryset)
                    is_list = True
                
                if filter_type == 'text':
                    queryset = [obj for obj in queryset if param_value.lower() in str(getattr(obj, field_name) or '').lower()]
                elif filter_type == 'select' or filter_type == 'date':
                    # Exact match
                    queryset = [obj for obj in queryset if str(getattr(obj, field_name)) == str(param_value)]
            else:
                # DB filtering
                if filter_type == 'text':
                    queryset = queryset.filter(**{f"{field_name}__icontains": param_value})
                elif filter_type == 'select' or filter_type == 'date':
                    queryset = queryset.filter(**{field_name: param_value})

        # Date Range Filter
        if config.get('type') == 'date':
            min_val = params.get(f"{field_name}_min")
            max_val = params.get(f"{field_name}_max")
            
            # Date range logic
            if min_val or max_val:
                if is_encrypted or is_list:
                    if not is_list:
                        queryset = list(queryset)
                        is_list = True
                    
                    if min_val:
                        queryset = [obj for obj in queryset if str(getattr(obj, field_name)) >= min_val]
                        config['value_min'] = min_val
                    if max_val:
                        queryset = [obj for obj in queryset if str(getattr(obj, field_name)) <= max_val]
                        config['value_max'] = max_val
                else:
                    if min_val:
                        queryset = queryset.filter(**{f"{field_name}__gte": min_val})
                        config['value_min'] = min_val
                    if max_val:
                        queryset = queryset.filter(**{f"{field_name}__lte": max_val})
                        config['value_max'] = max_val

    return queryset

def apply_sorting(queryset, sort_param):
    if not sort_param:
        return queryset

    # Check if we need Python sorting (list input or encrypted field)
    is_list = isinstance(queryset, list)
    model = queryset.model if hasattr(queryset, 'model') else (queryset[0].__class__ if queryset else None)
    
    is_encrypted_sort = False
    clean_sort_param = sort_param.lstrip('-')
    if model:
        try:
            field = model._meta.get_field(clean_sort_param)
            if isinstance(field, (EncryptedCharField, EncryptedTextField, EncryptedDateField, EncryptedBooleanField)):
                is_encrypted_sort = True
        except: pass

    if is_list or is_encrypted_sort:
        if not is_list:
            queryset = list(queryset)
        
        reverse = sort_param.startswith('-')
        try:
            return sorted(queryset, key=lambda x: getattr(x, clean_sort_param) or '', reverse=reverse)
        except:
            return queryset
    else:
        try:
            return queryset.order_by(sort_param)
        except:
            return queryset

# Generic helper for CRUD views
def generic_list_view(request, model_class, title, create_url_name, update_url_name, fields_to_show, bulk_delete_url_name=None, export_url_name=None, filter_config=None, import_url_name=None, queryset=None, extra_actions=None, mobile_config=None, extra_context=None):

    if queryset is not None:
        items = queryset
    else:
        items = model_class.objects.all()

    if filter_config:
        # Populate options for select fields dynamically if not provided
        for config in filter_config:
            if config['type'] == 'select' and 'options' not in config:
                # Assuming the field name refers to a related model (e.g., 'workplace')
                # Or a field with choices
                try:
                    field_object = model_class._meta.get_field(config['field'])
                    if field_object.is_relation:
                        related_model = field_object.related_model
                        config['options'] = [(obj.pk, str(obj)) for obj in related_model.objects.all()]
                    elif field_object.choices:
                        config['options'] = field_object.choices
                except:
                    config['options'] = []

        items = apply_filters(items, filter_config, request.GET)

    current_sort = request.GET.get('sort')
    items = apply_sorting(items, current_sort)

    context = {
        'items': items,
        'title': title,
        'create_url_name': create_url_name,
        'update_url_name': update_url_name,
        'bulk_delete_url_name': bulk_delete_url_name,
        'export_url_name': export_url_name,
        'import_url_name': import_url_name,
        'fields': fields_to_show,
        'filter_config': filter_config,
        'extra_actions': extra_actions,
        'mobile_config': mobile_config,
    }
    if extra_context:
        context.update(extra_context)

    return render(request, 'core/list_template.html', context)


def generic_import_view(request, model_class, title, list_url_name, step=1):
    handler = ImportHandler(request)

    # Step 1: Upload
    if step == 1:
        if request.method == 'POST':
            if 'import_file' in request.FILES:
                handler.save_file(request.FILES['import_file'])
                # Clear previous session data
                if 'import_mapping' in request.session: del request.session['import_mapping']
                if 'import_settings' in request.session: del request.session['import_settings']
                return redirect(f'import_{model_class.__name__.lower()}_step2')
        return render(request, 'core/import/import_step1.html', {'title': title, 'list_url_name': list_url_name})

    # Step 2: Settings
    elif step == 2:
        if request.method == 'POST':
            settings = {
                'delimiter': request.POST.get('delimiter', ';'),
                'date_format': request.POST.get('date_format', '%Y-%m-%d'),
                'encoding': request.POST.get('encoding', 'utf-8-sig'),
                'uppercase_names': request.POST.get('uppercase_names') == 'on'
            }
            request.session['import_settings'] = settings
            return redirect(f'import_{model_class.__name__.lower()}_step3')
        return render(request, 'core/import/import_step2.html', {'title': title})

    # Step 3: Mapping
    elif step == 3:
        settings = request.session.get('import_settings', {'delimiter': ';', 'encoding': 'utf-8-sig'})
        csv_headers = handler.get_headers(delimiter=settings['delimiter'], encoding=settings['encoding'])

        # Get Model Fields
        model_fields = []
        for field in model_class._meta.fields:
            if field.auto_created or field.name == 'id':
                continue

            field_type = field.get_internal_type()
            required = not field.blank and not field.null

            model_fields.append({
                'name': field.name,
                'verbose_name': field.verbose_name,
                'help_text': field.help_text,
                'required': required,
                'type': field_type
            })

        if request.method == 'POST':
            mapping = {}
            for field in model_fields:
                val = request.POST.get(f'map_{field["name"]}')
                mapping[field['name']] = val if val else None

            request.session['import_mapping'] = mapping
            return redirect(f'import_{model_class.__name__.lower()}_step4')

        return render(request, 'core/import/import_step3.html', {
            'title': title,
            'csv_headers': csv_headers,
            'model_fields': model_fields
        })

    # Step 4: Preview & Execute
    elif step == 4:
        settings = request.session.get('import_settings', {})
        mapping = request.session.get('import_mapping', {})

        if request.method == 'POST':
            count = handler.execute_import(
                model_class, mapping,
                delimiter=settings.get('delimiter', ';'),
                date_format=settings.get('date_format', '%Y-%m-%d'),
                encoding=settings.get('encoding', 'utf-8-sig')
            )
            messages.success(request, f'{count} kayıt başarıyla içe aktarıldı.')
            return redirect(list_url_name)

        summary = handler.get_preview_data(
            model_class, mapping,
            delimiter=settings.get('delimiter', ';'),
            date_format=settings.get('date_format', '%Y-%m-%d'),
            encoding=settings.get('encoding', 'utf-8-sig')
        )

        return render(request, 'core/import/import_step4.html', {'title': title, 'summary': summary})

    return redirect(list_url_name)

def generic_bulk_delete_view(request, model_class, list_url_name):
    if request.method == 'POST':
        selected_ids = request.POST.getlist('selected_items')
        if selected_ids:
            # Log deletion (need to fetch before delete to log)
            items = model_class.objects.filter(id__in=selected_ids)
            count = items.count()
            for item in items:
                log_action(request.user, 'Silme', item)

            items.delete()
            messages.success(request, f'{count} kayıt silindi.')
        else:
            messages.warning(request, 'Silinecek kayıt seçilmedi.')
    return redirect(list_url_name)

def generic_create_view(request, form_class, title, list_url_name, form_kwargs=None):
    if form_kwargs is None: form_kwargs = {}
    if request.method == 'POST':
        form = form_class(request.POST, **form_kwargs)
        if form.is_valid():
            obj = form.save()
            log_action(request.user, 'Oluşturma', obj)
            messages.success(request, 'Kayıt başarıyla oluşturuldu.')
            return redirect(list_url_name)
    else:
        form = form_class(**form_kwargs)
    return render(request, 'core/form_template.html', {'form': form, 'title': title})

def generic_update_view(request, model_class, form_class, pk, title, list_url_name, form_kwargs=None):
    if form_kwargs is None: form_kwargs = {}
    item = get_object_or_404(model_class, pk=pk)
    if request.method == 'POST':
        form = form_class(request.POST, instance=item, **form_kwargs)
        if form.is_valid():
            obj = form.save()
            log_action(request.user, 'Güncelleme', obj)
            messages.success(request, 'Kayıt güncellendi.')
            return redirect(list_url_name)
    else:
        form = form_class(instance=item, **form_kwargs)
    return render(request, 'core/form_template.html', {'form': form, 'title': title})

def generic_export_view(request, model_class, filter_config=None):
    queryset = model_class.objects.all()
    if filter_config:
         queryset = apply_filters(queryset, filter_config, request.GET)

    export_format = request.GET.get('format', 'csv')
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"{model_class._meta.verbose_name_plural}_{timestamp}"

    # Get all field names
    field_names = [field.name for field in model_class._meta.fields]
    # Make header pretty (verbose names)
    headers = [field.verbose_name for field in model_class._meta.fields]

    if export_format == 'xlsx':
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename="{filename}.xlsx"'

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Export"

        ws.append(headers)

        for obj in queryset:
            row = []
            for field in field_names:
                val = getattr(obj, field)
                # Handle relations nicely (str representation)
                if hasattr(val, 'pk'): # It's a related object
                    val = str(val)
                # Handle choices nicely
                elif hasattr(obj, f'get_{field}_display'):
                    val = getattr(obj, f'get_{field}_display')()

                # Handle dates and None
                if val is None:
                    val = ""
                # Convert dates to string for Excel compatibility if needed,
                # but openpyxl handles datetime objects well.
                # Just ensuring timezone info is stripped if present/problematic
                if isinstance(val, datetime):
                     val = val.replace(tzinfo=None)
                
                # Convert UUID to string for Excel compatibility
                from uuid import UUID
                if isinstance(val, UUID):
                    val = str(val)

                row.append(val)
            ws.append(row)

        wb.save(response)
        return response

    else: # Default to CSV
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="{filename}.csv"'
        response.write(u'\ufeff'.encode('utf8')) # BOM for Excel compatibility

        writer = csv.writer(response)
        writer.writerow(headers)

        for obj in queryset:
            row = []
            for field in field_names:
                val = getattr(obj, field)
                if hasattr(val, 'pk'):
                    val = str(val)
                elif hasattr(obj, f'get_{field}_display'):
                    val = getattr(obj, f'get_{field}_display')()
                if val is None:
                    val = ""
                row.append(val)
            writer.writerow(row)

        return response


# Specific Views using helpers
@login_required
def workplace_list(request):
    from django.db.models import Count, Q
    from datetime import date, timedelta

    # ACL: Centralized permission check
    queryset = get_allowed_workplaces(request.user)

    # Prefetch related data for cards
    queryset = queryset.prefetch_related(
        'workers__education_set', 
        'workers__examination_set',
        'facilities'
    )

    # Apply filters
    search_query = request.GET.get('search', '')
    hazard_filter = request.GET.get('hazard_class', '')
    
    # Apply DB filters first
    if hazard_filter:
        queryset = queryset.filter(hazard_class=hazard_filter)
        
    # Apply encrypted field search (Python side)
    if search_query:
        # Switch to list for Python filtering
        queryset = list(queryset)
        search_query_lower = search_query.lower()
        queryset = [
            obj for obj in queryset
            if (obj.name and search_query_lower in obj.name.lower()) or
               (obj.detsis_number and search_query_lower in obj.detsis_number.lower()) or
               (obj.nace_code and search_query_lower in obj.nace_code.lower())
        ]
    
    # Calculate summary statistics
    all_workplaces = Workplace.objects.all()
    total_count = all_workplaces.count()
    low_hazard_count = all_workplaces.filter(hazard_class='LOW').count()
    medium_hazard_count = all_workplaces.filter(hazard_class='MEDIUM').count()
    high_hazard_count = all_workplaces.filter(hazard_class='HIGH').count()
    
    # Calculate Stats (Scoped)
    stats = get_user_scoped_stats(request.user)

    context = {
        'items': queryset,
        'title': 'İşyerleri',
        'search_query': search_query,
        'hazard_filter': hazard_filter,
        'stats': stats,
        'create_url_name': 'workplace_create',
        'bulk_delete_url_name': 'workplace_bulk_delete',
        'import_url_name': 'import_workplace_step1',
        'export_url_name': 'workplace_export',
    }
    return render(request, 'core/workplace_list.html', context)
    
    if total_workers > 0:
        today = date.today()
        workers = Worker.objects.select_related('workplace').prefetch_related('education_set', 'examination_set')
        
        for worker in workers:
            wp = worker.workplace
            if wp:
                edu_years = wp.get_validity_years('education')
                exam_years = wp.get_validity_years('examination')
            else:
                edu_years = 2
                exam_years = 1
            
            # Valid education
            for edu in worker.education_set.all():
                if edu.date and edu.date >= today - relativedelta(years=edu_years):
                    valid_education += 1
                    break
            
            # Valid examination  
            for exam in worker.examination_set.all():
                if exam.date and exam.date >= today - relativedelta(years=exam_years):
                    valid_examination += 1
                    break
            
            # First aid count
            if worker.first_aid_certificate and worker.first_aid_expiry_date and worker.first_aid_expiry_date >= today:
                first_aid_count += 1
    
    education_percentage = round((valid_education / total_workers * 100) if total_workers > 0 else 0)
    examination_percentage = round((valid_examination / total_workers * 100) if total_workers > 0 else 0)

    context = {
        'items': queryset,
        'title': 'Firmalar',
        'search_query': search_query,
        'hazard_filter': hazard_filter,
        'total_count': total_count,
        'low_hazard_count': low_hazard_count,
        'medium_hazard_count': medium_hazard_count,
        'high_hazard_count': high_hazard_count,
        'total_facilities': total_facilities,
        'total_workers': total_workers,
        'education_percentage': education_percentage,
        'examination_percentage': examination_percentage,
        'first_aid_count': first_aid_count,
        'create_url_name': 'workplace_create',
        'update_url_name': 'workplace_update',
        'bulk_delete_url_name': 'workplace_bulk_delete',
        'import_url_name': 'import_workplace_step1',
        'export_url_name': 'workplace_export',
    }
    return render(request, 'core/workplace_list.html', context)


@login_required
def workplace_import(request, step=1):
    return generic_import_view(request, Workplace, "İşyeri İçe Aktar", 'workplace_list', step=step)

@login_required
def workplace_bulk_delete(request):
    return generic_bulk_delete_view(request, Workplace, 'workplace_list')

@login_required
def workplace_export(request):
    filter_config = [
        {'field': 'name', 'type': 'text'},
        {'field': 'detsis_number', 'type': 'text'},
        {'field': 'hazard_class', 'type': 'select'},
    ]
    return generic_export_view(request, Workplace, filter_config)

@login_required
def workplace_create(request):
    if request.method == 'POST':
        form = WorkplaceForm(request.POST)
        if form.is_valid():
            created_workplace = form.save()
            
            # Auto-create fallback facility
            Facility.objects.create(name="MERKEZ BİNA", workplace=created_workplace)
            
            log_action(request.user, 'Oluşturma', created_workplace)
            messages.success(request, 'İşyeri başarıyla oluşturuldu.')
            return redirect('workplace_detail', pk=created_workplace.pk)
    else:
        form = WorkplaceForm()

    return render(request, 'core/workplace_form.html', {
        'form': form,
        'title': "Yeni İş Yeri"
    })

@login_required
def workplace_detail(request, pk):
    from datetime import date
    from dateutil.relativedelta import relativedelta
    
    # Permission Check
    allowed = get_allowed_workplaces(request.user)
    workplace = get_object_or_404(allowed, pk=pk)
    
    # Get all facilities and workers
    facilities = workplace.facilities.prefetch_related('worker_set__education_set', 'worker_set__examination_set', 'assessment_sessions').all()
    workers = Worker.objects.filter(workplace=workplace).prefetch_related('education_set', 'examination_set')
    
    total_workers = workers.count()
    facility_count = facilities.count()
    
    # Calculate education percentage
    valid_education = 0
    valid_examination = 0
    first_aid_count = 0
    
    today = date.today()
    edu_years = workplace.get_validity_years('education')
    exam_years = workplace.get_validity_years('examination')
    
    for worker in workers:
        # Check valid education
        for edu in worker.education_set.all():
            if edu.date and edu.date >= today - relativedelta(years=edu_years):
                valid_education += 1
                break
        
        # Check valid examination
        for exam in worker.examination_set.all():
            if exam.date and exam.date >= today - relativedelta(years=exam_years):
                valid_examination += 1
                break
        
        # Count first aid certificates
        if worker.first_aid_certificate:
            if worker.first_aid_expiry_date and worker.first_aid_expiry_date >= today:
                first_aid_count += 1
    
    education_percentage = round((valid_education / total_workers * 100) if total_workers > 0 else 0)
    examination_percentage = round((valid_examination / total_workers * 100) if total_workers > 0 else 0)
    
    # Calculate per-facility compliance stats
    facilities_with_stats = []
    for facility in facilities:
        facility_workers = list(facility.worker_set.all())
        fac_worker_count = len(facility_workers)
        fac_valid_edu = 0
        fac_valid_exam = 0
        fac_first_aid = 0
        
        for worker in facility_workers:
            # Education
            for edu in worker.education_set.all():
                if edu.date and edu.date >= today - relativedelta(years=edu_years):
                    fac_valid_edu += 1
                    break
            # Examination
            for exam in worker.examination_set.all():
                if exam.date and exam.date >= today - relativedelta(years=exam_years):
                    fac_valid_exam += 1
                    break
            # First aid
            if worker.first_aid_certificate and worker.first_aid_expiry_date and worker.first_aid_expiry_date >= today:
                fac_first_aid += 1
        
        facilities_with_stats.append({
            'facility': facility,
            'worker_count': fac_worker_count,
            'edu_percentage': round((fac_valid_edu / fac_worker_count * 100) if fac_worker_count > 0 else 0),
            'exam_percentage': round((fac_valid_exam / fac_worker_count * 100) if fac_worker_count > 0 else 0),
            'first_aid_count': fac_first_aid,
            'assessment_count': facility.assessment_sessions.count(),
        })
    
    # Get tab from query param
    active_tab = request.GET.get('tab', 'overview')
    
    # Filter workers by facility if specified
    facility_filter = request.GET.get('facility', '')
    filtered_workers = workers
    if facility_filter:
        filtered_workers = workers.filter(facility_id=facility_filter)
    
    # Fetch education sessions for this workplace
    education_sessions = Education.objects.filter(
        workplace=workplace
    ).prefetch_related('workers', 'workers__facility', 'professionals').order_by('-date')
    
    context = {
        'workplace': workplace,
        'facilities': facilities,
        'facilities_with_stats': facilities_with_stats,
        'workers': filtered_workers,
        'total_workers': total_workers,
        'facility_count': facility_count,
        'valid_education': valid_education,
        'valid_examination': valid_examination,
        'education_percentage': education_percentage,
        'examination_percentage': examination_percentage,
        'first_aid_count': first_aid_count,
        'active_tab': active_tab,
        'facility_filter': facility_filter,
        'education_sessions': education_sessions,
    }
    
    return render(request, 'core/workplace_detail.html', context)

@login_required
def workplace_update(request, pk):
    allowed = get_allowed_workplaces(request.user)
    item = get_object_or_404(allowed, pk=pk)

    if request.method == 'POST':
        form = WorkplaceForm(request.POST, instance=item)
        if form.is_valid():
            form.save()
            log_action(request.user, 'Güncelleme', item)
            messages.success(request, 'Kayıt güncellendi.')
            return redirect('workplace_list')
    else:
        form = WorkplaceForm(instance=item)

    # Fetch facilities with their workers
    facilities = item.facilities.prefetch_related('worker_set').all()

    return render(request, 'core/workplace_form.html', {
        'form': form, 
        'workplace': item,
        'title': "İşyeri Düzenle", 
        'facilities': facilities
    })

@login_required
def worker_list(request):
    filter_config = [
        {'field': 'name', 'label': 'Ad Soyad', 'type': 'text'},
        {'field': 'tckn', 'label': 'TCKN', 'type': 'text'},
        {'field': 'workplace', 'label': 'İşyeri', 'type': 'select'},
        {'field': 'profession', 'label': 'Meslek', 'type': 'select'},
    ]

    # Prefetch related data to optimize badge generation
    allowed_workplaces = get_allowed_workplaces(request.user)
    # Fix: Use 'education' and 'examination' as per stats.py findings
    queryset = Worker.objects.filter(workplace__in=allowed_workplaces).select_related('workplace', 'facility').prefetch_related('education_set', 'examination_set')

    extra_actions = [
        {
            'url_name': 'examination_create',
            'label': 'Muayene Ekle',
            'icon': 'bi-heart-pulse',
            'btn_class': 'btn-outline-success',
            'query_param': 'worker_id'
        }
    ]

    # Calculate Stats for Context
    stats = get_user_scoped_stats(request.user)
    
    return generic_list_view(request, Worker, "Çalışanlar", 'worker_create', 'worker_update',
                             [('name', 'Ad Soyad'),
                              ('tckn', 'TCKN'),
                              ('workplace', 'İşyeri'),
                              ('facility', 'Bina/Birim'),
                              ('education_status', 'Eğitim Durumu'),
                              ('examination_status', 'Muayene Durumu')],
                             'worker_bulk_delete', 'worker_export', filter_config, 'import_worker_step1',
                             queryset=queryset, extra_actions=extra_actions, extra_context={'stats': stats})

@login_required
def worker_import(request, step=1):
    return generic_import_view(request, Worker, "Çalışan İçe Aktar", 'worker_list', step=step)

@login_required
def worker_bulk_delete(request):
    return generic_bulk_delete_view(request, Worker, 'worker_list')

@login_required
def worker_export(request):
    filter_config = [
        {'field': 'name', 'type': 'text'},
        {'field': 'tckn', 'type': 'text'},
        {'field': 'workplace', 'type': 'select'},
        {'field': 'profession', 'type': 'select'},
    ]
    return generic_export_view(request, Worker, filter_config)

@login_required
def worker_create(request):
    workplace_id = request.GET.get('workplace')
    facility_id = request.GET.get('facility')
    initial_data = {}
    
    if workplace_id:
        initial_data['workplace'] = workplace_id
    if facility_id:
        initial_data['facility'] = facility_id

    if request.method == 'POST':
        form = WorkerForm(request.POST, user=request.user)
        if form.is_valid():
            obj = form.save()
            log_action(request.user, 'Oluşturma', obj)
            messages.success(request, 'Kayıt başarıyla oluşturuldu.')
            if workplace_id:
                return redirect('workplace_detail', pk=workplace_id)
            return redirect('worker_list')
    else:
        form = WorkerForm(initial=initial_data, user=request.user)
        # Filter querysets if explicit ID provided (Double safety)
        if workplace_id:
            form.fields['workplace'].queryset = Workplace.objects.filter(pk=workplace_id)
        if facility_id:
            form.fields['facility'].queryset = Facility.objects.filter(pk=facility_id)
    
    return render(request, 'core/worker_form.html', {'form': form, 'title': "Yeni Çalışan"})

@login_required
def worker_update(request, pk):
    item = get_object_or_404(Worker, pk=pk)
    # Check permission? Ideally yes. But form filter also handles dropdowns.
    
    if request.method == 'POST':
        form = WorkerForm(request.POST, instance=item, user=request.user)
        if form.is_valid():
            obj = form.save()
            log_action(request.user, 'Güncelleme', obj)
            messages.success(request, 'Kayıt güncellendi.')
            return redirect('worker_list')
    else:
        form = WorkerForm(instance=item, user=request.user)

    examinations = item.examination_set.all().order_by('-date')
    return render(request, 'core/worker_form.html', {'form': form, 'title': "Çalışan Düzenle", 'examinations': examinations})

@login_required
def professional_list(request):
    filter_config = [
        {'field': 'name', 'label': 'Ad Soyad', 'type': 'text'},
        {'field': 'role', 'label': 'Görevi', 'type': 'select'},
    ]
    return generic_list_view(request, Professional, "Profesyoneller", 'professional_create', 'professional_update',
                             [('name', 'Ad Soyad'), ('license_id', 'Lisans No'), ('get_role_display', 'Görevi')],
                             'professional_bulk_delete', 'professional_export', filter_config, 'import_professional_step1')

@login_required
def professional_import(request, step=1):
    return generic_import_view(request, Professional, "Profesyonel İçe Aktar", 'professional_list', step=step)

@login_required
def professional_bulk_delete(request):
    return generic_bulk_delete_view(request, Professional, 'professional_list')

@login_required
def professional_export(request):
    filter_config = [
        {'field': 'name', 'type': 'text'},
        {'field': 'role', 'type': 'select'},
    ]
    return generic_export_view(request, Professional, filter_config)

@login_required
def professional_create(request):
    return generic_create_view(request, ProfessionalForm, "Yeni Profesyonel", 'professional_list')

@login_required
def professional_update(request, pk):
    return generic_update_view(request, Professional, ProfessionalForm, pk, "Profesyonel Düzenle", 'professional_list')

@login_required
def education_list(request):
    filter_config = [
        {'field': 'topic', 'label': 'Konu', 'type': 'text'},
        {'field': 'date', 'label': 'Tarih', 'type': 'date'},
        {'field': 'workplace', 'label': 'İşyeri', 'type': 'select'},
        {'field': 'professionals', 'label': 'Eğiticiler', 'type': 'select'},
    ]
    return generic_list_view(request, Education, "İSG Eğitimleri", 'education_create', 'education_update',
                             [('date', 'Tarih'), ('topic', 'Konu'), ('workplace', 'İşyeri')],
                             'education_bulk_delete', 'education_export', filter_config, 'import_education_step1',
                             extra_actions=[{
                                 'url_name': 'education_certificate_download',
                                 'label': 'PDF',
                                 'icon': 'bi-file-pdf',
                                 'btn_class': 'btn-outline-danger',
                                 'query_param': 'education_id'
                             }, {
                                 'url_name': 'education_certificate_word',
                                 'label': 'Word',
                                 'icon': 'bi-file-word',
                                 'btn_class': 'btn-outline-primary',
                                 'query_param': 'education_id'
                             }])

@login_required
def education_certificate_download(request):
    education_id = request.GET.get('education_id')
    education = get_object_or_404(Education, pk=education_id)

    pdf_buffer = generate_certificate_pdf(education)

    filename = f"Sertifika_{education.id}_{datetime.now().strftime('%Y%m%d')}.pdf"
    return FileResponse(pdf_buffer, as_attachment=True, filename=filename)


@login_required
def education_participation_form(request, pk):
    """Export participation form PDF for an education session."""
    from .pdf_generator import generate_participation_form_pdf
    
    education = get_object_or_404(Education, pk=pk)
    pdf_buffer = generate_participation_form_pdf(education)
    
    filename = f"Katılım_Formu_{education.id}_{datetime.now().strftime('%Y%m%d')}.pdf"
    return FileResponse(pdf_buffer, as_attachment=True, filename=filename)

@login_required
def education_import(request, step=1):
    return generic_import_view(request, Education, "Eğitim İçe Aktar", 'education_list', step=step)

@login_required
def education_bulk_delete(request):
    return generic_bulk_delete_view(request, Education, 'education_list')

@login_required
def education_export(request):
    filter_config = [
        {'field': 'topic', 'type': 'text'},
        {'field': 'date', 'type': 'date'},
        {'field': 'workplace', 'type': 'select'},
    ]
    return generic_export_view(request, Education, filter_config)

@login_required
def education_create(request):
    from .models import EducationSubjects, EducationTopic
    
    workplace_id = request.GET.get('workplace')
    initial_data = {'duration': 8}  # Default duration to 8 hours
    has_workplace = False
    
    if workplace_id:
        initial_data['workplace'] = workplace_id
        has_workplace = True
        # Get all workers for this workplace to pre-select them
        workers = Worker.objects.filter(workplace_id=workplace_id)
        initial_data['workers'] = workers  # Pre-select all workers
    
    if request.method == 'POST':
        form = EducationForm(request.POST, user=request.user)
        if form.is_valid():
            education = form.save()
            
            # Save education topics
            topic_codes = request.POST.getlist('topic_codes[]')
            for code in topic_codes:
                duration = request.POST.get(f'topic_duration_{code}', 30)
                try:
                    duration = int(duration) if duration else 30
                except ValueError:
                    duration = 30
                EducationTopic.objects.create(
                    education=education,
                    topic_code=code,
                    duration_minutes=duration
                )
            
            messages.success(request, 'Kayıt başarıyla oluşturuldu.')
            # Redirect back to workplace detail if came from there
            if workplace_id:
                return redirect('workplace_detail', pk=workplace_id)
            return redirect('education_list')
    else:
        form = EducationForm(initial=initial_data, user=request.user)
        # Filter querysets based on workplace if ID provided
        if workplace_id:
            form.fields['workplace'].queryset = Workplace.objects.filter(pk=workplace_id)
            form.fields['workers'].queryset = Worker.objects.filter(workplace_id=workplace_id)
    
    return render(request, 'core/education_form.html', {
        'form': form, 
        'title': "Yeni Eğitim",
        'has_workplace': has_workplace,
        'education_subjects': EducationSubjects.choices,
        'selected_topics': [code for code, label in EducationSubjects.choices],
        'topic_durations': json.dumps({}),
    })

@login_required
def education_update(request, pk):
    from .models import EducationSubjects, EducationTopic
    
    item = get_object_or_404(Education, pk=pk)
    
    # Get existing topics for this education
    existing_topics = {t.topic_code: t.duration_minutes for t in item.education_topics.all()}
    
    if request.method == 'POST':
        form = EducationForm(request.POST, instance=item, user=request.user)
        if form.is_valid():
            education = form.save()
            
            # Clear existing topics and re-create
            education.education_topics.all().delete()
            
            # Save education topics
            topic_codes = request.POST.getlist('topic_codes[]')
            for code in topic_codes:
                duration = request.POST.get(f'topic_duration_{code}', 30)
                try:
                    duration = int(duration) if duration else 30
                except ValueError:
                    duration = 30
                EducationTopic.objects.create(
                    education=education,
                    topic_code=code,
                    duration_minutes=duration
                )
            
            messages.success(request, 'Kayıt güncellendi.')
            return redirect('education_list')
    else:
        form = EducationForm(instance=item, user=request.user)
    
    return render(request, 'core/education_form.html', {
        'form': form, 
        'title': "Eğitim Düzenle",
        'education': item,
        'education_subjects': EducationSubjects.choices,
        'selected_topics': list(existing_topics.keys()),
        'topic_durations': json.dumps(existing_topics),
    })

@login_required
def inspection_list(request):
    filter_config = [
        {'field': 'date', 'label': 'Tarih', 'type': 'date'},
        {'field': 'workplace', 'label': 'İşyeri', 'type': 'select'},
        {'field': 'professional', 'label': 'Denetleyen', 'type': 'select'},
    ]
    return generic_list_view(request, Inspection, "Denetimler", 'inspection_create', 'inspection_update',
                             [('date', 'Tarih'), ('workplace', 'İşyeri'), ('notes', 'Notlar')],
                             'inspection_bulk_delete', 'inspection_export', filter_config, 'import_inspection_step1')

@login_required
def inspection_import(request, step=1):
    return generic_import_view(request, Inspection, "Denetim İçe Aktar", 'inspection_list', step=step)

@login_required
def inspection_bulk_delete(request):
    return generic_bulk_delete_view(request, Inspection, 'inspection_list')

@login_required
def inspection_export(request):
    filter_config = [
        {'field': 'date', 'type': 'date'},
        {'field': 'workplace', 'type': 'select'},
        {'field': 'professional', 'type': 'select'},
    ]
    return generic_export_view(request, Inspection, filter_config)

@login_required
def inspection_create(request):
    return generic_create_view(request, InspectionForm, "Yeni Denetim", 'inspection_list', form_kwargs={'user': request.user})

@login_required
def inspection_update(request, pk):
    return generic_update_view(request, Inspection, InspectionForm, pk, "Denetim Düzenle", 'inspection_list', form_kwargs={'user': request.user})

@login_required
@medical_access_required
def examination_list(request):
    filter_config = [
        {'field': 'date', 'label': 'Tarih', 'type': 'date'},
        {'field': 'worker__workplace', 'label': 'İşyeri', 'type': 'select'},
        {'field': 'professional', 'label': 'Hekim', 'type': 'select'},
        {'field': 'decision', 'label': 'Karar', 'type': 'select'},
    ]
    return generic_list_view(request, Examination, "Sağlık Muayeneleri", 'examination_create', 'examination_update',
                             [('caution_icon_html', ''), ('date', 'Tarih'), ('worker', 'Çalışan'), ('get_decision_display', 'Karar')],
                             'examination_bulk_delete', 'examination_export', filter_config, 'import_examination_step1')

@login_required
@medical_access_required
def examination_import(request, step=1):
    return generic_import_view(request, Examination, "Muayene İçe Aktar", 'examination_list', step=step)

@login_required
@medical_access_required
def examination_bulk_delete(request):
    return generic_bulk_delete_view(request, Examination, 'examination_list')

@login_required
@medical_access_required
def examination_export(request):
    filter_config = [
        {'field': 'date', 'type': 'date'},
        {'field': 'worker', 'type': 'select'},
        {'field': 'professional', 'type': 'select'},
        {'field': 'decision', 'type': 'select'},
    ]
    return generic_export_view(request, Examination, filter_config)

@login_required
@medical_access_required
def examination_create(request):
    if request.method == 'POST':
        form = ExaminationForm(request.POST, user=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, 'Kayıt başarıyla oluşturuldu.')
            return redirect('examination_list')
    else:
        initial_data = {}
        worker_id = request.GET.get('worker_id')
        if worker_id:
            worker = get_object_or_404(Worker, pk=worker_id)
            initial_data['worker'] = worker
        form = ExaminationForm(initial=initial_data, user=request.user)
    return render(request, 'core/examination_form.html', {'form': form, 'title': "Yeni Muayene"})

@login_required
@medical_access_required
def examination_update(request, pk):
    item = get_object_or_404(Examination, pk=pk)
    if request.method == 'POST':
        form = ExaminationForm(request.POST, instance=item, user=request.user)
        if form.is_valid():
            form.save()
            messages.success(request, 'Kayıt güncellendi.')
            return redirect('examination_list')
    else:
        form = ExaminationForm(instance=item, user=request.user)
    return render(request, 'core/examination_form.html', {'form': form, 'title': "Muayene Düzenle"})

@login_required
@medical_access_required
def update_examination_note(request, pk):
    examination = get_object_or_404(Examination, pk=pk)
    if request.method == 'POST':
        action = request.POST.get('action')
        if action == 'delete':
            examination.caution_note = ''
            examination.is_caution = False
            examination.save()
            messages.success(request, 'Uyarı notu silindi.')
        else:
            form = ExaminationNoteForm(request.POST, instance=examination)
            if form.is_valid():
                examination.is_caution = True # Ensure flag is set if note is saved
                form.save()
                messages.success(request, 'Uyarı notu güncellendi.')
            else:
                messages.error(request, 'Hata oluştu.')
    return redirect('examination_list')

# Facility Views
@login_required
def facility_list(request):
    from django.db.models import Count, Q
    from datetime import date
    from dateutil.relativedelta import relativedelta

    try:
        # Base Query
        allowed_workplaces = get_allowed_workplaces(request.user)
        queryset = Facility.objects.filter(workplace__in=allowed_workplaces)\
            .select_related('workplace').prefetch_related(
                'worker_set__education_set',
                'worker_set__examination_set'
            ).annotate(total_workers_count=Count('worker'))

        # Search
        search_query = request.GET.get('search', '')
        if search_query:
            queryset = queryset.filter(Q(name__icontains=search_query) | Q(workplace__name__icontains=search_query))

        # Filter by Workplace
        workplace_filter = request.GET.get('workplace')
        if workplace_filter:
            try:
                workplace_filter = int(workplace_filter)
                queryset = queryset.filter(workplace_id=workplace_filter)
            except ValueError:
                workplace_filter = None

        # Stats Calculation - Dynamic based on filter
        # If workplace filter is active, calculate stats for that workplace only
        if workplace_filter:
            total_facilities = queryset.count()
            # Get workers from filtered facilities
            workers_qs = Worker.objects.filter(
                facility__workplace_id=workplace_filter
            ).select_related('workplace').prefetch_related('education_set', 'examination_set')
        else:
            total_facilities = Facility.objects.count()
            workers_qs = Worker.objects.select_related('workplace').prefetch_related('education_set', 'examination_set')
        
        # Calculate Stats
        stats = get_user_scoped_stats(request.user)

        context = {
            'items': queryset,
            'title': 'Birimler',
            'search_query': search_query,
            'workplace_filter': workplace_filter,
            'workplaces': allowed_workplaces,
            'stats': stats,
            'create_url_name': 'facility_create',
            'bulk_delete_url_name': 'facility_bulk_delete',
            'import_url_name': 'import_facility_step1',
            'export_url_name': 'facility_export',
        }
        return render(request, 'core/facility_list.html', context)
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise e

@login_required
def facility_bulk_delete(request):
    return generic_bulk_delete_view(request, Facility, 'facility_list')

@login_required
def facility_create(request):
    workplace_id = request.GET.get('workplace')
    initial_data = {}
    
    if workplace_id:
        initial_data['workplace'] = workplace_id
    
    if request.method == 'POST':
        form = FacilityForm(request.POST, user=request.user)
        if form.is_valid():
            obj = form.save()
            log_action(request.user, 'Oluşturma', obj)
            messages.success(request, 'Kayıt başarıyla oluşturuldu.')
            # Redirect back to workplace detail if came from there
            if workplace_id:
                return redirect('workplace_detail', pk=workplace_id)
            return redirect('facility_list')
    else:
        form = FacilityForm(initial=initial_data, user=request.user)
        # Filter workplace queryset if workplace_id provided
        if workplace_id:
            form.fields['workplace'].queryset = Workplace.objects.filter(pk=workplace_id)
    
    return render(request, 'core/form_template.html', {'form': form, 'title': "Yeni Bina/Birim"})

@login_required
def facility_detail(request, pk):
    """Facility Dashboard - Read-only overview with tabs"""
    allowed_workplaces = get_allowed_workplaces(request.user)
    facility = get_object_or_404(Facility, pk=pk, workplace__in=allowed_workplaces)
    workplace = facility.workplace
    
    # Workers for this facility
    workers = facility.worker_set.all()
    worker_count = workers.count()
    
    # Calculate compliance stats
    from datetime import date
    from dateutil.relativedelta import relativedelta
    
    valid_education = 0
    valid_examination = 0
    today = date.today()
    
    if worker_count > 0:
        # Get validity periods from workplace
        edu_years = workplace.get_validity_years('education')
        exam_years = workplace.get_validity_years('examination')
        
        for worker in workers:
            # Check education validity
            for edu in worker.education_set.all():
                if edu.date and edu.date >= today - relativedelta(years=edu_years):
                    valid_education += 1
                    break
            
            # Check examination validity
            for exam in worker.examination_set.all():
                if exam.date and exam.date >= today - relativedelta(years=exam_years):
                    valid_examination += 1
                    break
    
    education_percentage = round((valid_education / worker_count * 100) if worker_count > 0 else 0)
    examination_percentage = round((valid_examination / worker_count * 100) if worker_count > 0 else 0)
    
    # Risk assessments for this facility
    assessment_sessions = facility.assessment_sessions.all().order_by('-created_at')
    assessment_count = assessment_sessions.count()
    
    # Available risk tools for new assessment
    risk_tools = RiskTool.objects.filter(is_active=True)
    
    # Default session form for the modal
    default_title = f"Risk Değerlendirmesi - {today.strftime('%d.%m.%Y')}"
    session_form = AssessmentSessionForm(initial={'title': default_title})
    
    # Parse and format coordinates for map (4 decimal places)
    lat, lng = None, None
    lat_display, lng_display = None, None
    if facility.coordinates:
        try:
            parts = facility.coordinates.replace(' ', '').split(',')
            if len(parts) == 2:
                lat, lng = float(parts[0]), float(parts[1])
                lat_display = round(lat, 4)
                lng_display = round(lng, 4)
        except (ValueError, TypeError):
            pass
    
    context = {
        'facility': facility,
        'workplace': workplace,
        'workers': workers,
        'worker_count': worker_count,
        'valid_education': valid_education,
        'valid_examination': valid_examination,
        'education_percentage': education_percentage,
        'examination_percentage': examination_percentage,
        'assessment_sessions': assessment_sessions,
        'assessment_count': assessment_count,
        'risk_tools': risk_tools,
        'session_form': session_form,
        'lat': lat,
        'lng': lng,
        'lat_display': lat_display,
        'lng_display': lng_display,
        'title': facility.name,
    }
    return render(request, 'core/facility_detail.html', context)


@login_required
def facility_update(request, pk):
    allowed_workplaces = get_allowed_workplaces(request.user)
    item = get_object_or_404(Facility, pk=pk, workplace__in=allowed_workplaces)
    if request.method == 'POST':
        form = FacilityForm(request.POST, instance=item, user=request.user)
        if form.is_valid():
            obj = form.save()
            log_action(request.user, 'Güncelleme', obj)
            messages.success(request, 'Kayıt güncellendi.')
            return redirect('facility_detail', pk=pk)
    else:
        form = FacilityForm(instance=item, user=request.user)
    
    return render(request, 'core/facility_settings.html', {
        'form': form,
        'title': "Birim Ayarları",
        'facility': item,
    })

@login_required
def facility_import(request, step=1):
    return generic_import_view(request, Facility, "Bina/Birim İçe Aktar", 'facility_list', step=step)

@login_required
def facility_export(request):
    filter_config = [
        {'field': 'name', 'type': 'text'},
        {'field': 'workplace', 'type': 'select'},
    ]
    return generic_export_view(request, Facility, filter_config)

@login_required
def settings_view(request):
    return render(request, 'core/settings.html', {'title': 'Ayarlar'})

@login_required
def certificate_settings_view(request):
    template, created = CertificateTemplate.objects.get_or_create(name="Global")

    if request.method == 'POST':
        form = CertificateTemplateForm(request.POST, instance=template)
        if form.is_valid():
            form.save()
            messages.success(request, 'Ayarlar kaydedildi.')
            return redirect('certificate_settings')
    else:
        form = CertificateTemplateForm(instance=template)

    return render(request, 'core/certificate_settings.html', {
        'title': 'Sertifika Tasarımı',
        'form': form
    })

@login_required
def user_list(request):
    # Fetch all users
    # Use prefetch_related for reverse OneToOne (safer) or select_related if Django supports it (it does but can be tricky)
    # Removing select_related('profile') to avoid potential "Invalid field name" error if it fails on reverse
    # We will handle it in the loop or use prefetch_related
    users_qs = User.objects.prefetch_related('assignments__workplace').order_by('username')
    
    # Search
    search_query = request.GET.get('search', '')

    if search_query:
        users_qs = users_qs.filter(username__icontains=search_query)

    # Convert to list to persist attributes and avoid template re-query
    users = list(users_qs)

    # Safely attach profile for template
    for user in users:
        try:
            # Try to access profile. if prefetch didn't happen, it queries DB.
            # If DoesNotExist, we catch it.
            user.safe_profile = user.profile
        except Exception: # Catch UserProfile.DoesNotExist and any other DB issue
            user.safe_profile = None

    # Safely attach profile for Request User (for Sidebar)
    try:
        if hasattr(request.user, 'profile'):
            request.user.safe_profile = request.user.profile
        else:
            request.user.safe_profile = None
    except Exception:
        request.user.safe_profile = None

    workplaces = Workplace.objects.all()

    # Calculate Stats
    total_users = len(users)
    active_users = sum(1 for u in users if u.is_active)
    
    # Role counts (safe_profile might be None)
    specialist_count = sum(1 for u in users if u.safe_profile and u.safe_profile.role == 'SPECIALIST')
    doctor_count = sum(1 for u in users if u.safe_profile and u.safe_profile.role == 'DOCTOR')

    # Serialize assignments for JS
    # Structure: { user_id: [ { workplace_name, start_date, end_date, is_active } ] }
    user_assignments = {}
    for user in users:
        # Skip admins/managers/superusers from having assignments in this context if desired, or include all.
        # Logic in template was: if not user.is_superuser and role != 'ADMIN' and role != 'MANAGER'
        # We replicate that here or just send all. Let's replicate for consistency.
        role = user.safe_profile.role if user.safe_profile else ''
        if not user.is_superuser and role != 'ADMIN' and role != 'MANAGER':
            assigns = []
            for assign in user.assignments.all():
                assigns.append({
                    'id': assign.pk,
                    'workplace_name': assign.workplace.name,
                    'start_date': assign.start_date.strftime('%Y-%m-%d'),
                    'end_date': assign.end_date.strftime('%Y-%m-%d') if assign.end_date else '',
                    'is_active': assign.is_active
                })
            user_assignments[user.pk] = assigns

    context = {
        'users': users,
        'workplaces': workplaces,
        'title': 'Kullanıcılar',
        'total_users': total_users,
        'active_users': active_users,
        'specialist_count': specialist_count,
        'doctor_count': doctor_count,
        'user_assignments': user_assignments,
    }
    return render(request, 'core/user_list.html', context)

@login_required
@require_POST
def user_update_profile_ajax(request, pk):
    """Update user role and details via AJAX (KVKK Audit Logged)"""
    user = get_object_or_404(User, pk=pk)
    
    # Check permissions (Only Admin or Manager)
    if not request.user.is_superuser:
        if not hasattr(request.user, 'profile') or request.user.profile.role not in ['ADMIN', 'MANAGER']:
            return JsonResponse({'error': 'Unauthorized'}, status=403)

    try:
        data = json.loads(request.body)
        profile, created = UserProfile.objects.get_or_create(user=user)
        
        # 1. Update Role
        new_role = data.get('role')
        if new_role and profile.role != new_role:
            # SAFETY CHECK: If current user is ADMIN, check if they are the last one
            if profile.role == 'ADMIN':
                admin_count = UserProfile.objects.filter(role='ADMIN').count()
                if admin_count <= 1:
                    return JsonResponse({'error': 'Sistemde en az bir Sistem Yöneticisi kalmalıdır. Son yönetici rolünü değiştiremezsiniz.'}, status=400)

            # Audit Log
            ActionLog.objects.create(
                user=request.user,
                action='ROLE_CHANGE',
                model_name='UserProfile',
                object_id=str(user.pk),
                details=f"Role changed from {profile.role} to {new_role}"
            )
            profile.role = new_role

        # 2. Update TCKN/Phone/MFA (if provided)
        # Note: In a real app we'd validate these strictly.
        if 'tckn' in data: profile.tckn = data['tckn']
        if 'phone' in data: profile.phone = data['phone']
        if 'is_mfa_enabled' in data: profile.is_mfa_enabled = data['is_mfa_enabled']
        if 'is_active' in data: 
            user.is_active = data['is_active']
            user.save()

        profile.save()
        return JsonResponse({'status': 'success'})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
@require_POST
def api_delete_assignment(request, pk):
    """Delete a workplace assignment"""
    assignment = get_object_or_404(WorkplaceAssignment, pk=pk)
    
    # Permissions: Admin, or Manager
    if not request.user.is_superuser:
         if not hasattr(request.user, 'profile') or request.user.profile.role not in ['ADMIN', 'MANAGER']:
            return JsonResponse({'error': 'Unauthorized'}, status=403)

    try:
        # Log before delete
        ActionLog.objects.create(
            user=request.user,
            action='DELETE_ASSIGNMENT',
            model_name='WorkplaceAssignment',
            object_id=str(assignment.pk),
            details=f"Deleted assignment: User {assignment.user.username} - {assignment.workplace.name}"
        )
        assignment.delete()
        return JsonResponse({'status': 'success'})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
@require_POST
def api_update_assignment(request, pk):
    """Update an existing workplace assignment"""
    assignment = get_object_or_404(WorkplaceAssignment, pk=pk)

    # Permissions: Admin, or Manager
    if not request.user.is_superuser:
         if not hasattr(request.user, 'profile') or request.user.profile.role not in ['ADMIN', 'MANAGER']:
            return JsonResponse({'error': 'Unauthorized'}, status=403)

    try:
        data = json.loads(request.body)
        start_date = data.get('start_date')
        end_date = data.get('end_date')

        if not start_date:
            return JsonResponse({'error': 'Start date is required'}, status=400)

        # Log before update
        ActionLog.objects.create(
            user=request.user,
            action='UPDATE_ASSIGNMENT',
            model_name='WorkplaceAssignment',
            object_id=str(assignment.pk),
            details=f"Updated assignment {assignment.pk}: dates changed."
        )

        assignment.start_date = start_date
        assignment.end_date = end_date if end_date else None
        assignment.save()
        
        return JsonResponse({'status': 'success'})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
@require_POST
def workplace_assignment_create(request):
    """Create a new workplace assignment (validity period)"""
    try:
        data = json.loads(request.body)
        user_id = data.get('user_id')
        workplace_id = data.get('workplace_id')
        start_date = data.get('start_date')
        end_date = data.get('end_date')
        role = data.get('role', '')

        if not user_id or not workplace_id or not start_date:
            return JsonResponse({'error': 'Missing required fields'}, status=400)

        assignment = WorkplaceAssignment.objects.create(
            user_id=user_id,
            workplace_id=workplace_id,
            role=role,
            start_date=start_date,
            end_date=end_date if end_date else None
        )
        
        # Audit Log
        ActionLog.objects.create(
            user=request.user,
            action='ASSIGNMENT_CREATE',
            model_name='WorkplaceAssignment',
            object_id=str(assignment.pk),
            details=f"Assigned user {user_id} to workplace {workplace_id}"
        )
        
        return JsonResponse({'status': 'success'})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
@require_POST
def workplace_assignment_revoke(request, pk):
    """Revoke (deactivate) an assignment"""
    try:
        assignment = get_object_or_404(WorkplaceAssignment, pk=pk)
        assignment.is_active = False
        assignment.save()
        
        # Audit Log
        ActionLog.objects.create(
            user=request.user,
            action='ASSIGNMENT_REVOKE',
            model_name='WorkplaceAssignment',
            object_id=str(assignment.pk),
            details=f"Revoked assignment {assignment.pk}"
        )
        
        return JsonResponse({'status': 'success'})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)

@login_required
def user_create(request):
    return generic_create_view(request, CustomUserCreationForm, "Yeni Kullanıcı", 'user_list')

@login_required
def user_update(request, pk):
    return generic_update_view(request, User, CustomUserChangeForm, pk, "Kullanıcı Düzenle", 'user_list')

@login_required
def user_bulk_delete(request):
    return generic_bulk_delete_view(request, User, 'user_list')

@login_required
def log_list(request):
    filter_config = [
        {'field': 'user', 'label': 'Kullanıcı', 'type': 'select'},
        {'field': 'action', 'label': 'İşlem', 'type': 'text'},
        {'field': 'model_name', 'label': 'Veri Türü', 'type': 'text'},
    ]
    # For logs, we typically want read-only, so we might need to adjust generic_list_view or pass empty create/update urls
    # But generic_list_view expects url names. We can pass placeholders or None if we modify generic view to handle None.
    # The current generic_list_view renders update button if url is provided.
    # Let's pass None for update_url_name to disable edit button.
    # We need to update generic_list_view template to check for None update_url_name.

    queryset = ActionLog.objects.all()

    # We need to handle 'user' filter which is a relationship. generic view's apply_filters handles select for foreign keys.
    # But ActionLog user is standard User model. apply_filters tries to find options from related model.
    # It should work.

    return generic_list_view(request, ActionLog, "İşlem Kayıtları", None, None,
                             [('timestamp', 'Zaman'), ('user', 'Kullanıcı'), ('action', 'İşlem'), ('model_name', 'Veri Türü'), ('details', 'Detaylar')],
                             None, None, filter_config, None, queryset=queryset)

@login_required
def profession_list(request):
    filter_config = [{'field': 'name', 'label': 'Meslek Adı', 'type': 'text'}]
    return generic_list_view(request, Profession, "Meslekler", 'profession_create', 'profession_update',
                             [('name', 'Meslek Adı')],
                             'profession_bulk_delete', 'profession_export', filter_config, 'import_profession_step1')

@login_required
def profession_bulk_delete(request):
    return generic_bulk_delete_view(request, Profession, 'profession_list')

@login_required
def profession_export(request):
    return generic_export_view(request, Profession, [{'field': 'name', 'type': 'text'}])

@login_required
def profession_import(request, step=1):
    return generic_import_view(request, Profession, "Meslek İçe Aktar", 'profession_list', step=step)

@login_required
def profession_create(request):
    return generic_create_view(request, ProfessionForm, "Yeni Meslek", 'profession_list')

@login_required
def profession_update(request, pk):
    return generic_update_view(request, Profession, ProfessionForm, pk, "Meslek Düzenle", 'profession_list')

# Statistics Configuration
STATISTICS_CONFIG = {
    'Workplace': {
        'fields': [
            {'name': 'hazard_class', 'label': 'Tehlike Sınıfı', 'type': 'category'},
        ]
    },
    'Worker': {
        'fields': [
            {'name': 'workplace__name', 'label': 'İşyeri', 'type': 'category'},
            {'name': 'gender', 'label': 'Cinsiyet', 'type': 'category'},
            {'name': 'profession__name', 'label': 'Meslek', 'type': 'category'},
        ]
    },
    'Examination': {
        'fields': [
            {'name': 'decision', 'label': 'Karar', 'type': 'category'},
            {'name': 'worker__workplace__name', 'label': 'İşyeri', 'type': 'category'},
            {'name': 'professional__name', 'label': 'Hekim', 'type': 'category'},
            {'name': 'date', 'label': 'Tarih', 'type': 'date'},
        ]
    },
    'Education': {
        'fields': [
            {'name': 'topic', 'label': 'Konu', 'type': 'category'},
            {'name': 'professionals__name', 'label': 'Eğiticiler', 'type': 'category'},
            {'name': 'workplace__name', 'label': 'İşyeri', 'type': 'category'},
            {'name': 'date', 'label': 'Tarih', 'type': 'date'},
        ]
    }
}

@login_required
def statistics_view(request):
    return render(request, 'core/statistics.html', {'config': STATISTICS_CONFIG})

@login_required
def api_get_statistics(request):
    model_name = request.GET.get('model')
    x_field = request.GET.get('x_field')
    y_field = request.GET.get('y_field') # Optional Grouping

    if not model_name or not x_field:
        return JsonResponse({'error': 'Missing parameters'}, status=400)

    # Resolve Model
    model_map = {
        'Workplace': Workplace,
        'Worker': Worker,
        'Examination': Examination,
        'Education': Education
    }
    model_class = model_map.get(model_name)
    if not model_class:
        return JsonResponse({'error': 'Invalid model'}, status=400)

    allowed_workplaces = get_allowed_workplaces(request.user)
    if model_name == 'Workplace':
        queryset = allowed_workplaces
    elif model_name == 'Worker':
        queryset = Worker.objects.filter(workplace__in=allowed_workplaces)
    elif model_name == 'Education':
        queryset = Education.objects.filter(workplace__in=allowed_workplaces)
    elif model_name == 'Examination':
        queryset = Examination.objects.filter(worker__workplace__in=allowed_workplaces)
    else:
        # Fallback for safety
        queryset = model_class.objects.none()
    
    queryset = queryset.distinct()

    # Aggregation
    try:
        from django.db.models import Count
        if y_field:
            data = list(queryset.values(x_field, y_field).annotate(count=Count('id', distinct=True)).order_by(x_field, y_field))
        else:
            data = list(queryset.values(x_field).annotate(count=Count('id', distinct=True)).order_by(x_field))

        # Helper to find field object
        def get_field_choices(model, field_path):
            parts = field_path.split('__')
            current_model = model
            field = None
            for part in parts:
                try:
                    field = current_model._meta.get_field(part)
                    if field.is_relation and field.related_model:
                        current_model = field.related_model
                except: return None
            return field.choices if field else None

        x_choices = get_field_choices(model_class, x_field)
        y_choices = get_field_choices(model_class, y_field) if y_field else None

        processed_data = []
        for row in data:
            x_val = row[x_field]
            y_val = row.get(y_field)

            # Map X
            if x_choices:
                for k, v in x_choices:
                    if str(k) == str(x_val):
                        x_val = v; break

            # Map Y
            if y_field and y_choices:
                for k, v in y_choices:
                    if str(k) == str(y_val):
                        y_val = v; break

            # Handle None
            if x_val is None: x_val = "Belirtilmemiş"
            if y_field and y_val is None: y_val = "Belirtilmemiş"

            item = {'x': x_val, 'count': row['count']}
            if y_field:
                item['y'] = y_val
            processed_data.append(item)

        return JsonResponse({'data': processed_data})

    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


# =============================================================================
# Risk Assessment Views
# =============================================================================

@login_required
def assessment_session_create(request, facility_id):
    """Create a new assessment session for a facility (called from modal)"""
    facility = get_object_or_404(Facility, pk=facility_id)
    
    if request.method == 'POST':
        from datetime import date
        workflow_type = request.POST.get('workflow_type', 'LIBRARY')
        scoring_method = request.POST.get('scoring_method', 'KINNEY')
        tool_id = request.POST.get('tool') if workflow_type == 'TEMPLATE' else None
        title = request.POST.get('title', f"Risk Değerlendirmesi - {date.today().strftime('%d.%m.%Y')}")
        
        # Create session
        session = AssessmentSession.objects.create(
            facility=facility,
            tool_id=tool_id if tool_id else None,
            title=title,
            workflow_type=workflow_type,
            scoring_method=scoring_method,
        )
        
        log_action(request.user, 'Oluşturma', session)
        messages.success(request, 'Değerlendirme başarıyla oluşturuldu.')
        
        # Redirect based on workflow type
        # Both workflows go to team page first for team member selection
        return redirect('assessment_team', session_pk=session.pk)
    
    # If GET or form invalid, redirect back to facility
    return redirect('facility_detail', pk=facility_id)


@login_required
def assessment_session_run(request, pk):
    """OiRA-style assessment runner with split-screen layout"""
    session = get_object_or_404(AssessmentSession, pk=pk)
    
    # Check workflow_type: LIBRARY assessments go to fast-track runner
    workflow_type = session.workflow_type or 'TEMPLATE'  # Default to TEMPLATE if NULL
    if workflow_type == 'LIBRARY':
        return redirect('assessment_fast_run', pk=session.pk)
    
    tool = session.tool
    
    # If no tool is set, redirect to facility (shouldn't happen for TEMPLATE workflow)
    if not tool:
        messages.warning(request, 'Bu değerlendirme için araç seçilmemiş.')
        return redirect('facility_update', pk=session.facility.pk)
    
    # Get all questions for this tool (ordered by category, topic, then question)
    all_questions = RiskQuestion.objects.filter(
        topic__category__tool=tool
    ).select_related('topic__category').order_by(
        'topic__category__order_index',
        'topic__order_index',
        'order_index'
    )
    
    questions_list = list(all_questions)
    total_questions = len(questions_list)
    
    if total_questions == 0:
        messages.warning(request, 'Bu değerlendirme aracında henüz soru bulunmuyor.')
        return redirect('facility_update', pk=session.facility.pk)
    
    # Get existing answers for this session
    existing_answers = {
        answer.question_id: answer 
        for answer in session.answers.all()
    }
    
    # Build navigation structure with answer status
    categories_data = []
    current_category = None
    current_topic = None
    
    for idx, question in enumerate(questions_list):
        cat = question.topic.category
        topic = question.topic
        
        # New category?
        if current_category is None or current_category['id'] != cat.id:
            current_category = {
                'id': cat.id,
                'title': cat.title,
                'topics': []
            }
            categories_data.append(current_category)
            current_topic = None
        
        # New topic?
        if current_topic is None or current_topic['id'] != topic.id:
            current_topic = {
                'id': topic.id,
                'title': topic.title,
                'questions': []
            }
            current_category['topics'].append(current_topic)
        
        # Get answer status
        answer = existing_answers.get(question.id)
        status = 'unanswered'
        if answer and answer.response:
            if answer.response == 'YES':
                status = 'yes'
            elif answer.response == 'NO':
                status = 'no'
            elif answer.response == 'NA':
                status = 'na'
            else:
                status = 'postponed'
        
        current_topic['questions'].append({
            'id': question.id,
            'index': idx,
            'content': question.content[:60] + '...' if len(question.content) > 60 else question.content,
            'status': status
        })
    
    # Determine current question (from ?q= param or first unanswered)
    q_param = request.GET.get('q')
    current_question = None
    current_index = 0
    
    if q_param:
        try:
            q_id = int(q_param)
            for idx, q in enumerate(questions_list):
                if q.id == q_id:
                    current_question = q
                    current_index = idx
                    break
        except (ValueError, TypeError):
            pass
    
    # Default to first unanswered or first question
    if not current_question:
        for idx, q in enumerate(questions_list):
            if q.id not in existing_answers or not existing_answers[q.id].response:
                current_question = q
                current_index = idx
                break
        if not current_question:
            current_question = questions_list[0]
            current_index = 0
    
    # Get current answer if exists
    current_answer = existing_answers.get(current_question.id)
    
    # Calculate prev/next
    prev_question = questions_list[current_index - 1] if current_index > 0 else None
    next_question = questions_list[current_index + 1] if current_index < total_questions - 1 else None
    
    # Get custom risks for sidebar
    custom_risks = session.custom_risks.all()
    
    context = {
        'session': session,
        'facility': session.facility,
        'tool': tool,
        'title': session.title,
        'categories_data': categories_data,
        'current_question': current_question,
        'current_answer': current_answer,
        'current_index': current_index,
        'total_questions': total_questions,
        'prev_question': prev_question,
        'next_question': next_question,
        'progress_percentage': session.progress_percentage,
        'custom_risks': custom_risks,
        'custom_risks_count': custom_risks.count(),
    }
    return render(request, 'core/assessment_run.html', context)


# =============================================================================
# Risk Tool Management Views
# =============================================================================

@login_required
def risk_tool_list(request):
    """List all Risk Assessment Tools"""
    tools = RiskTool.objects.all()
    context = {
        'tools': tools,
        'title': 'Değerlendirme Formları',
    }
    return render(request, 'core/risk_tool_list.html', context)


@login_required
def risk_tool_create(request):
    """Create a new Risk Tool manually"""
    if request.method == 'POST':
        form = RiskToolForm(request.POST)
        if form.is_valid():
            tool = form.save()
            log_action(request.user, 'Oluşturma', tool)
            messages.success(request, 'Risk değerlendirme aracı oluşturuldu.')
            return redirect('risk_tool_list')
    else:
        form = RiskToolForm()
    
    return render(request, 'core/risk_tool_form.html', {
        'form': form,
        'title': 'Yeni Değerlendirme Aracı'
    })


@login_required
def risk_tool_update(request, pk):
    """Edit an existing Risk Tool"""
    tool = get_object_or_404(RiskTool, pk=pk)
    if request.method == 'POST':
        form = RiskToolForm(request.POST, instance=tool)
        if form.is_valid():
            form.save()
            log_action(request.user, 'Güncelleme', tool)
            messages.success(request, 'Risk değerlendirme aracı güncellendi.')
            return redirect('risk_tool_list')
    else:
        form = RiskToolForm(instance=tool)
    
    return render(request, 'core/risk_tool_form.html', {
        'form': form,
        'title': 'Değerlendirme Aracı Düzenle',
        'tool': tool
    })


@login_required
def risk_tool_delete(request, pk):
    """Delete a Risk Tool"""
    tool = get_object_or_404(RiskTool, pk=pk)
    if request.method == 'POST':
        tool_title = tool.title
        tool.delete()
        messages.success(request, f'"{tool_title}" silindi.')
        return redirect('risk_tool_list')
    
    return render(request, 'core/confirm_delete.html', {
        'object': tool,
        'title': 'Değerlendirme Aracını Sil',
        'cancel_url': 'risk_tool_list'
    })


@login_required
def risk_tool_import(request):
    """Import Risk Tool from Excel/CSV file"""
    import pandas as pd
    import os
    
    if request.method == 'POST':
        form = RiskToolImportForm(request.POST, request.FILES)
        if form.is_valid():
            uploaded_file = request.FILES['file']
            tool_name = form.cleaned_data.get('tool_name')
            sector = form.cleaned_data.get('sector', '')
            
            # Use filename if tool_name not provided
            if not tool_name:
                tool_name = os.path.splitext(uploaded_file.name)[0]
            
            try:
                # Read file based on extension
                if uploaded_file.name.endswith('.csv'):
                    df = pd.read_csv(uploaded_file)
                else:
                    df = pd.read_excel(uploaded_file)
                
                # Normalize column names (lowercase, strip whitespace)
                df.columns = df.columns.str.lower().str.strip()
                
                # Validate required columns
                required_cols = ['category', 'topic', 'question']
                missing_cols = [col for col in required_cols if col not in df.columns]
                if missing_cols:
                    messages.error(request, f"Eksik sütunlar: {', '.join(missing_cols)}")
                    return redirect('risk_tool_import')
                
                # Create the RiskTool
                tool = RiskTool.objects.create(
                    title=tool_name,
                    sector=sector,
                    is_active=True
                )
                
                # Track created objects
                categories_created = 0
                topics_created = 0
                questions_created = 0
                
                # Process each row
                category_cache = {}
                topic_cache = {}
                
                for idx, row in df.iterrows():
                    category_name = str(row.get('category', '')).strip()
                    topic_name = str(row.get('topic', '')).strip()
                    question_text = str(row.get('question', '')).strip()
                    explanation = str(row.get('explanation', '')).strip() if 'explanation' in df.columns else ''
                    legal_ref = str(row.get('legal_ref', '')).strip() if 'legal_ref' in df.columns else ''
                    
                    if not category_name or not topic_name or not question_text:
                        continue
                    
                    # Get or create Category
                    if category_name not in category_cache:
                        category, created = RiskCategory.objects.get_or_create(
                            tool=tool,
                            title=category_name,
                            defaults={'order_index': len(category_cache)}
                        )
                        category_cache[category_name] = category
                        if created:
                            categories_created += 1
                    category = category_cache[category_name]
                    
                    # Get or create Topic
                    topic_key = f"{category_name}|{topic_name}"
                    if topic_key not in topic_cache:
                        topic, created = RiskTopic.objects.get_or_create(
                            category=category,
                            title=topic_name,
                            defaults={'order_index': len([k for k in topic_cache if k.startswith(category_name)])}
                        )
                        topic_cache[topic_key] = topic
                        if created:
                            topics_created += 1
                    topic = topic_cache[topic_key]
                    
                    # Create Question
                    RiskQuestion.objects.create(
                        topic=topic,
                        content=question_text,
                        explanation_text=explanation,
                        legal_reference=legal_ref,
                        order_index=questions_created
                    )
                    questions_created += 1
                
                log_action(request.user, 'İçe Aktarma', tool, f"{questions_created} soru eklendi")
                messages.success(
                    request,
                    f'"{tool_name}" başarıyla içe aktarıldı! '
                    f'{categories_created} kategori, {topics_created} konu, {questions_created} soru oluşturuldu.'
                )
                return redirect('risk_tool_list')
                
            except Exception as e:
                messages.error(request, f'Dosya işlenirken hata oluştu: {str(e)}')
                return redirect('risk_tool_import')
    else:
        form = RiskToolImportForm()
    
    return render(request, 'core/risk_tool_import.html', {
        'form': form,
        'title': 'Excel/CSV İçe Aktar'
    })


@login_required
def risk_tool_template_download(request):
    """Download a blank CSV template"""
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="risk_degerlendirme_sablonu.csv"'
    response.write('\ufeff')  # UTF-8 BOM for Excel compatibility
    
    writer = csv.writer(response)
    writer.writerow(['Category', 'Topic', 'Question', 'Explanation', 'Legal_Ref'])
    writer.writerow(['Fiziksel Riskler', 'Gürültü', 'İşyerinde gürültü ölçümü yapılıyor mu?', 'Yıllık gürültü ölçümü yapılmalıdır.', ''])
    writer.writerow(['Fiziksel Riskler', 'Aydınlatma', 'Çalışma alanlarında yeterli aydınlatma var mı?', '', ''])
    writer.writerow(['Ergonomik Riskler', 'Oturma Düzeni', 'Sandalyeler ergonomik mi?', '', ''])
    
    return response


@login_required
def assessment_answer_save(request, session_pk):
    """Save an answer via AJAX POST"""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    
    session = get_object_or_404(AssessmentSession, pk=session_pk)
    
    try:
        data = json.loads(request.body)
        question_id = data.get('question_id')
        response_value = data.get('response')
        notes = data.get('notes', '')
        
        if not question_id or not response_value:
            return JsonResponse({'error': 'Missing data'}, status=400)
        
        question = get_object_or_404(RiskQuestion, pk=question_id)
        
        # Get or create answer
        from .models import AssessmentAnswer
        answer, created = AssessmentAnswer.objects.get_or_create(
            session=session,
            question=question,
            defaults={'response': response_value, 'notes': notes}
        )
        
        if not created:
            answer.response = response_value
            answer.notes = notes
            answer.save()
        
        # Update session timestamp
        session.save()  # This triggers auto_now on updated_at
        
        return JsonResponse({
            'success': True,
            'created': created,
            'progress': session.progress_percentage
        })
        
    except (json.JSONDecodeError, KeyError) as e:
        return JsonResponse({'error': str(e)}, status=400)


# =============================================================================
# Custom Risks Views
# =============================================================================

def get_runner_context(session):
    """Build context for runner sidebar - shared by all runner views"""
    tool = session.tool
    
    # Get all questions for this tool
    all_questions = RiskQuestion.objects.filter(
        topic__category__tool=tool
    ).select_related('topic__category').order_by(
        'topic__category__order_index',
        'topic__order_index',
        'order_index'
    )
    
    questions_list = list(all_questions)
    total_questions = len(questions_list)
    
    # Get existing answers
    existing_answers = {
        answer.question_id: answer 
        for answer in session.answers.all()
    }
    
    # Build navigation structure
    categories_data = []
    current_category = None
    current_topic = None
    
    for idx, question in enumerate(questions_list):
        cat = question.topic.category
        topic = question.topic
        
        if current_category is None or current_category['id'] != cat.id:
            current_category = {'id': cat.id, 'title': cat.title, 'topics': []}
            categories_data.append(current_category)
            current_topic = None
        
        if current_topic is None or current_topic['id'] != topic.id:
            current_topic = {'id': topic.id, 'title': topic.title, 'questions': []}
            current_category['topics'].append(current_topic)
        
        answer = existing_answers.get(question.id)
        status = 'unanswered'
        if answer and answer.response:
            if answer.response == 'YES':
                status = 'yes'
            elif answer.response == 'NO':
                status = 'no'
            elif answer.response == 'NA':
                status = 'na'
        
        current_topic['questions'].append({
            'id': question.id,
            'index': idx,
            'content': question.content[:60] + '...' if len(question.content) > 60 else question.content,
            'status': status
        })
    
    custom_risks = session.custom_risks.all()
    
    return {
        'session': session,
        'facility': session.facility,
        'tool': tool,
        'categories_data': categories_data,
        'total_questions': total_questions,
        'progress_percentage': session.progress_percentage,
        'custom_risks': custom_risks,
        'custom_risks_count': custom_risks.count(),
    }


@login_required
def custom_risk_list(request, session_pk):
    """List custom risks with runner layout"""
    session = get_object_or_404(AssessmentSession, pk=session_pk)
    context = get_runner_context(session)
    context['view_mode'] = 'custom_list'
    return render(request, 'core/custom_risk_runner.html', context)


@login_required
def custom_risk_create(request, session_pk):
    """Create a new custom risk with runner layout"""
    session = get_object_or_404(AssessmentSession, pk=session_pk)
    
    if request.method == 'POST':
        form = CustomRiskForm(request.POST)
        if form.is_valid():
            custom_risk = form.save(commit=False)
            custom_risk.session = session
            custom_risk.save()
            messages.success(request, 'Ek risk eklendi.')
            
            if 'save_another' in request.POST:
                return redirect('custom_risk_create', session_pk=session.pk)
            return redirect('custom_risk_list', session_pk=session.pk)
    else:
        form = CustomRiskForm()
    
    context = get_runner_context(session)
    context['form'] = form
    context['view_mode'] = 'custom_add'
    context['title'] = 'Yeni Ek Risk'
    return render(request, 'core/custom_risk_runner.html', context)


@login_required
def custom_risk_update(request, session_pk, pk):
    """Edit custom risk with runner layout"""
    session = get_object_or_404(AssessmentSession, pk=session_pk)
    custom_risk = get_object_or_404(AssessmentCustomRisk, pk=pk, session=session)
    
    if request.method == 'POST':
        form = CustomRiskForm(request.POST, instance=custom_risk)
        if form.is_valid():
            form.save()
            messages.success(request, 'Ek risk güncellendi.')
            return redirect('custom_risk_list', session_pk=session.pk)
    else:
        form = CustomRiskForm(instance=custom_risk)
    
    context = get_runner_context(session)
    context['form'] = form
    context['custom_risk'] = custom_risk
    context['view_mode'] = 'custom_edit'
    context['title'] = 'Ek Riski Düzenle'
    return render(request, 'core/custom_risk_runner.html', context)


@login_required
def custom_risk_delete(request, session_pk, pk):
    """Delete a custom risk"""
    session = get_object_or_404(AssessmentSession, pk=session_pk)
    custom_risk = get_object_or_404(AssessmentCustomRisk, pk=pk, session=session)
    
    if request.method == 'POST':
        custom_risk.delete()
        messages.success(request, 'Ek risk silindi.')
        return redirect('custom_risk_list', session_pk=session.pk)
    
    context = get_runner_context(session)
    context['object'] = custom_risk
    context['title'] = 'Ek Riski Sil'
    context['view_mode'] = 'custom_delete'
    return render(request, 'core/custom_risk_runner.html', context)


# =============================================================================
# Action Plan Views
# =============================================================================

def get_action_plan_context(session):
    """Build context for action plan sidebar - only shows 'No' risks"""
    tool = session.tool
    
    # Get answers with NO response (risks that need action)
    no_answers = AssessmentAnswer.objects.filter(
        session=session,
        response='NO'
    ).select_related('question__topic__category').order_by(
        'question__topic__category__order_index',
        'question__topic__order_index',
        'question__order_index'
    )
    
    # Build sidebar data for standard risks
    standard_risks = []
    for answer in no_answers:
        standard_risks.append({
            'id': answer.id,
            'type': 'standard',
            'title': answer.question.content[:60] + '...' if len(answer.question.content) > 60 else answer.question.content,
            'full_title': answer.question.content,
            'category': answer.question.topic.category.title,
            'topic': answer.question.topic.title,
            'priority': answer.risk_priority,
            'status': answer.action_plan_status,
            'measures_count': answer.measures.count(),
        })
    
    # Get custom risks with is_acceptable=False
    custom_risks_no = session.custom_risks.filter(is_acceptable=False)
    custom_risks_list = []
    for cr in custom_risks_no:
        custom_risks_list.append({
            'id': cr.id,
            'type': 'custom',
            'title': cr.description[:60] + '...' if len(cr.description) > 60 else cr.description,
            'full_title': cr.description,
            'priority': cr.risk_priority,
            'status': cr.action_plan_status,
            'measures_count': cr.custom_measures.count(),
        })
    
    return {
        'session': session,
        'facility': session.facility,
        'tool': tool,
        'standard_risks': standard_risks,
        'custom_risks_no': custom_risks_list,
        'total_risks': len(standard_risks) + len(custom_risks_list),
    }


@login_required
def action_plan_intro(request, session_pk):
    """Action plan landing/intro page"""
    session = get_object_or_404(AssessmentSession, pk=session_pk)
    context = get_action_plan_context(session)
    context['view_mode'] = 'intro'
    return render(request, 'core/action_plan_runner.html', context)


@login_required
def action_plan_list(request, session_pk):
    """List all risks requiring action"""
    session = get_object_or_404(AssessmentSession, pk=session_pk)
    context = get_action_plan_context(session)
    context['view_mode'] = 'list'
    return render(request, 'core/action_plan_runner.html', context)


@login_required
def action_plan_edit(request, session_pk, risk_type, risk_id):
    """Edit action plan for a specific risk"""
    session = get_object_or_404(AssessmentSession, pk=session_pk)
    
    if risk_type == 'standard':
        answer = get_object_or_404(AssessmentAnswer, pk=risk_id, session=session)
        risk_title = answer.question.content
        risk_info = answer.question.explanation_text or ''
        measures = answer.measures.all()
        priority = answer.risk_priority
        risk_obj = answer
        original_notes = answer.notes
    else:  # custom
        custom_risk = get_object_or_404(AssessmentCustomRisk, pk=risk_id, session=session)
        risk_title = custom_risk.description
        risk_info = custom_risk.evidence or ''
        measures = custom_risk.custom_measures.all()
        priority = custom_risk.risk_priority
        risk_obj = custom_risk
        original_notes = custom_risk.notes
    
    context = get_action_plan_context(session)
    
    # Build prev/next navigation
    all_risks = context['standard_risks'] + context['custom_risks_no']
    current_idx = None
    for idx, r in enumerate(all_risks):
        if r['type'] == risk_type and r['id'] == risk_id:
            current_idx = idx
            break
    
    prev_risk = all_risks[current_idx - 1] if current_idx and current_idx > 0 else None
    next_risk = all_risks[current_idx + 1] if current_idx is not None and current_idx < len(all_risks) - 1 else None
    
    context['view_mode'] = 'edit'
    context['risk_type'] = risk_type
    context['risk_id'] = risk_id
    context['risk_title'] = risk_title
    context['risk_info'] = risk_info
    context['measures'] = measures
    context['priority'] = priority
    context['risk_obj'] = risk_obj
    context['original_notes'] = original_notes
    context['current_risk'] = {'type': risk_type, 'id': risk_id}
    context['prev_risk'] = prev_risk
    context['next_risk'] = next_risk
    context['risk_index'] = current_idx + 1 if current_idx is not None else 1
    return render(request, 'core/action_plan_runner.html', context)


@login_required
def action_plan_priority_update(request, session_pk, risk_type, risk_id):
    """Update risk priority via AJAX"""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    
    session = get_object_or_404(AssessmentSession, pk=session_pk)
    
    try:
        data = json.loads(request.body)
        priority = data.get('priority')
        
        if risk_type == 'standard':
            answer = get_object_or_404(AssessmentAnswer, pk=risk_id, session=session)
            answer.risk_priority = priority
            answer.save()
        else:
            custom_risk = get_object_or_404(AssessmentCustomRisk, pk=risk_id, session=session)
            custom_risk.risk_priority = priority
            custom_risk.save()
        
        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=400)


@login_required
def measure_add(request, session_pk, risk_type, risk_id):
    """Add a new measure to a risk"""
    session = get_object_or_404(AssessmentSession, pk=session_pk)
    
    if request.method == 'POST':
        measure = ActionPlanMeasure()
        
        if risk_type == 'standard':
            answer = get_object_or_404(AssessmentAnswer, pk=risk_id, session=session)
            measure.answer = answer
        else:
            custom_risk = get_object_or_404(AssessmentCustomRisk, pk=risk_id, session=session)
            measure.custom_risk = custom_risk
        
        measure.description = request.POST.get('description', '')
        measure.expertise = request.POST.get('expertise', '')
        measure.responsible_person = request.POST.get('responsible_person', '')
        measure.budget = request.POST.get('budget', '')
        
        start_date = request.POST.get('planning_start_date')
        end_date = request.POST.get('planning_end_date')
        if start_date:
            measure.planning_start_date = start_date
        if end_date:
            measure.planning_end_date = end_date
        
        measure.save()
        messages.success(request, 'Önlem eklendi.')
    
    return redirect('action_plan_edit', session_pk=session_pk, risk_type=risk_type, risk_id=risk_id)


@login_required
def measure_update(request, session_pk, measure_id):
    """Update an existing measure"""
    session = get_object_or_404(AssessmentSession, pk=session_pk)
    measure = get_object_or_404(ActionPlanMeasure, pk=measure_id)
    
    # Determine risk type and id for redirect
    if measure.answer:
        risk_type = 'standard'
        risk_id = measure.answer.id
    else:
        risk_type = 'custom'
        risk_id = measure.custom_risk.id
    
    if request.method == 'POST':
        measure.description = request.POST.get('description', '')
        measure.expertise = request.POST.get('expertise', '')
        measure.responsible_person = request.POST.get('responsible_person', '')
        measure.budget = request.POST.get('budget', '')
        
        start_date = request.POST.get('planning_start_date')
        end_date = request.POST.get('planning_end_date')
        measure.planning_start_date = start_date if start_date else None
        measure.planning_end_date = end_date if end_date else None
        
        measure.save()
        messages.success(request, 'Önlem güncellendi.')
    
    return redirect('action_plan_edit', session_pk=session_pk, risk_type=risk_type, risk_id=risk_id)


@login_required
def measure_delete(request, session_pk, measure_id):
    """Delete a measure"""
    session = get_object_or_404(AssessmentSession, pk=session_pk)
    measure = get_object_or_404(ActionPlanMeasure, pk=measure_id)
    
    # Determine risk type and id for redirect
    if measure.answer:
        risk_type = 'standard'
        risk_id = measure.answer.id
    else:
        risk_type = 'custom'
        risk_id = measure.custom_risk.id
    
    if request.method == 'POST':
        measure.delete()
        messages.success(request, 'Önlem silindi.')
    
    return redirect('action_plan_edit', session_pk=session_pk, risk_type=risk_type, risk_id=risk_id)


# =============================================================================
# Status Dashboard & Reporting
# =============================================================================

@login_required
def assessment_status(request, session_pk):
    """Status dashboard showing assessment progress"""
    session = get_object_or_404(AssessmentSession, pk=session_pk)
    tool = session.tool
    
    # Get all categories with progress
    categories = RiskCategory.objects.filter(tool=tool).prefetch_related('topics__questions')
    
    category_progress = []
    total_questions = 0
    total_answered = 0
    
    for cat in categories:
        cat_questions = 0
        cat_answered = 0
        for topic in cat.topics.all():
            for q in topic.questions.all():
                cat_questions += 1
                total_questions += 1
                answer = session.answers.filter(question=q).first()
                if answer and answer.response:
                    cat_answered += 1
                    total_answered += 1
        
        progress = round((cat_answered / cat_questions * 100)) if cat_questions > 0 else 0
        category_progress.append({
            'title': cat.title,
            'answered': cat_answered,
            'total': cat_questions,
            'progress': progress,
        })
    
    # Risk summary
    high_priority = session.answers.filter(response='NO', risk_priority='HIGH').count()
    high_priority += session.custom_risks.filter(is_acceptable=False, risk_priority='HIGH').count()
    
    medium_priority = session.answers.filter(response='NO', risk_priority='MEDIUM').count()
    medium_priority += session.custom_risks.filter(is_acceptable=False, risk_priority='MEDIUM').count()
    
    low_priority = session.answers.filter(response='NO', risk_priority='LOW').count()
    low_priority += session.custom_risks.filter(is_acceptable=False, risk_priority='LOW').count()
    
    # Measures defined
    total_risks = session.answers.filter(response='NO').count() + session.custom_risks.filter(is_acceptable=False).count()
    risks_with_measures = 0
    for answer in session.answers.filter(response='NO'):
        if answer.measures.exists():
            risks_with_measures += 1
    for cr in session.custom_risks.filter(is_acceptable=False):
        if cr.custom_measures.exists():
            risks_with_measures += 1
    
    context = {
        'session': session,
        'facility': session.facility,
        'tool': tool,
        'category_progress': category_progress,
        'total_progress': round((total_answered / total_questions * 100)) if total_questions > 0 else 0,
        'high_priority': high_priority,
        'medium_priority': medium_priority,
        'low_priority': low_priority,
        'total_risks': total_risks,
        'risks_with_measures': risks_with_measures,
    }
    return render(request, 'core/assessment_status.html', context)


@login_required
def assessment_report(request, session_pk):
    """Report generation page"""
    session = get_object_or_404(AssessmentSession, pk=session_pk)
    
    if request.method == 'POST':
        # Save final comments and participants
        session.final_comments = request.POST.get('final_comments', '')
        session.participants = request.POST.get('participants', '')
        session.status = 'COMPLETED'
        session.save()
        messages.success(request, 'Değerlendirme tamamlandı!')
        return redirect('assessment_status', session_pk=session.pk)
    
    # Get team members for this session
    team_members = session.team_members.all()
    
    # Check if this is a library-based assessment
    is_library_workflow = session.workflow_type == 'LIBRARY'
    
    context = {
        'session': session,
        'facility': session.facility,
        'tool': session.tool,
        'team_members': team_members,
        'is_library_workflow': is_library_workflow,
    }
    return render(request, 'core/assessment_report.html', context)


@login_required
def export_action_plan_excel(request, session_pk):
    """Export action plan as Excel file"""
    session = get_object_or_404(AssessmentSession, pk=session_pk)
    
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Eylem Planı"
    
    # Headers
    headers = ['Risk No', 'Risk Tanımı', 'Kategori', 'Öncelik', 'Önlem', 'Sorumlu', 'Bütçe', 'Başlangıç', 'Bitiş']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = openpyxl.styles.Font(bold=True)
    
    row = 2
    risk_num = 1
    
    # Standard risks
    for answer in session.answers.filter(response='NO').select_related('question__topic__category'):
        measures = answer.measures.all()
        if measures:
            for measure in measures:
                ws.cell(row=row, column=1, value=f"R{risk_num}")
                ws.cell(row=row, column=2, value=answer.question.content)
                ws.cell(row=row, column=3, value=answer.question.topic.category.title)
                ws.cell(row=row, column=4, value=answer.get_risk_priority_display() if answer.risk_priority else '-')
                ws.cell(row=row, column=5, value=measure.description)
                ws.cell(row=row, column=6, value=measure.responsible_person)
                ws.cell(row=row, column=7, value=measure.budget)
                ws.cell(row=row, column=8, value=str(measure.planning_start_date) if measure.planning_start_date else '')
                ws.cell(row=row, column=9, value=str(measure.planning_end_date) if measure.planning_end_date else '')
                row += 1
        else:
            ws.cell(row=row, column=1, value=f"R{risk_num}")
            ws.cell(row=row, column=2, value=answer.question.content)
            ws.cell(row=row, column=3, value=answer.question.topic.category.title)
            ws.cell(row=row, column=4, value=answer.get_risk_priority_display() if answer.risk_priority else '-')
            ws.cell(row=row, column=5, value='Önlem tanımlanmadı')
            row += 1
        risk_num += 1
    
    # Custom risks
    for cr in session.custom_risks.filter(is_acceptable=False):
        measures = cr.custom_measures.all()
        if measures:
            for measure in measures:
                ws.cell(row=row, column=1, value=f"Ω{risk_num}")
                ws.cell(row=row, column=2, value=cr.description)
                ws.cell(row=row, column=3, value='Ek Risk')
                ws.cell(row=row, column=4, value=cr.get_risk_priority_display() if cr.risk_priority else '-')
                ws.cell(row=row, column=5, value=measure.description)
                ws.cell(row=row, column=6, value=measure.responsible_person)
                ws.cell(row=row, column=7, value=measure.budget)
                ws.cell(row=row, column=8, value=str(measure.planning_start_date) if measure.planning_start_date else '')
                ws.cell(row=row, column=9, value=str(measure.planning_end_date) if measure.planning_end_date else '')
                row += 1
        else:
            ws.cell(row=row, column=1, value=f"Ω{risk_num}")
            ws.cell(row=row, column=2, value=cr.description)
            ws.cell(row=row, column=3, value='Ek Risk')
            ws.cell(row=row, column=4, value=cr.get_risk_priority_display() if cr.risk_priority else '-')
            ws.cell(row=row, column=5, value='Önlem tanımlanmadı')
            row += 1
        risk_num += 1
    
    # Auto-size columns
    for col in ws.columns:
        max_length = max(len(str(cell.value or '')) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_length + 2, 50)
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="eylem_plani_{session.pk}.xlsx"'
    wb.save(response)
    return response


@login_required
def export_report_word(request, session_pk):
    """Export full report as Word document"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return HttpResponse("Error: python-docx library is not installed. Please run: pip install python-docx", status=500)
    
    try:
        session = get_object_or_404(AssessmentSession, pk=session_pk)
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Risk Değerlendirme Raporu', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Info
        doc.add_paragraph(f"Tesis: {session.facility.name}")
        doc.add_paragraph(f"Değerlendirme Aracı: {session.tool.title}")
        doc.add_paragraph(f"Tarih: {session.created_at.strftime('%d.%m.%Y')}")
        
        # Participants
        if session.participants:
            doc.add_paragraph(f"Katılımcılar: {session.participants}")
        
        doc.add_paragraph()
        
        # Methodology Section
        doc.add_heading('Metodoloji', level=1)
        methodology = f"""Bu değerlendirme, "{session.tool.title}" metodolojisi kullanılarak gerçekleştirilmiştir. 
Değerlendirme, 6331 sayılı İş Sağlığı ve Güvenliği Kanunu ve ilgili yönetmelikler çerçevesinde tehlikeleri tespit etmek, 
riskleri değerlendirmek ve gerekli önlemleri tanımlamak amacıyla yapılmıştır.

Değerlendirme kapsamında işyerindeki tüm faaliyetler, çalışma alanları ve süreçler incelenmiştir. 
Tespit edilen riskler öncelik sırasına göre sınıflandırılmış ve her biri için uygun kontrol önlemleri belirlenmiştir."""
        doc.add_paragraph(methodology)
        
        # Integrated Checklist (Denetim Listesi)
        doc.add_heading('Denetim Listesi', level=1)
        doc.add_paragraph("Tüm değerlendirme soruları ve cevapları:")
        
        # Get all questions for this tool
        questions = RiskQuestion.objects.filter(topic__category__tool=session.tool).select_related('topic__category')
        
        # Create checklist table
        checklist_table = doc.add_table(rows=1, cols=3)
        checklist_table.style = 'Table Grid'
        header_cells = checklist_table.rows[0].cells
        header_cells[0].text = 'Kategori'
        header_cells[1].text = 'Soru'
        header_cells[2].text = 'Durum'
        
        # Set header style
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
                    run.font.bold = True
        
        # Add questions
        for q in questions:
            answer = session.answers.filter(question=q).first()
            if answer and answer.response:
                if answer.response == 'YES':
                    status = '✔ Evet'
                elif answer.response == 'NO':
                    status = '❌ Hayır'
                elif answer.response == 'NA':
                    status = '— İlgili Değil'
                else:
                    status = answer.get_response_display()
            else:
                status = '⚪ Boş'
            
            row_cells = checklist_table.add_row().cells
            row_cells[0].text = q.topic.category.title[:30]
            row_cells[1].text = q.content[:100]
            row_cells[2].text = status
            
            # Small font for table rows
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
        
        doc.add_paragraph()
        
        # Executive Summary
        if session.final_comments:
            doc.add_heading('Yönetici Özeti', level=1)
            doc.add_paragraph(session.final_comments)
        
        # Risk Identification
        doc.add_heading('Tespit Edilen Riskler', level=1)
        
        risk_num = 1
        for answer in session.answers.filter(response='NO').select_related('question__topic__category'):
            doc.add_heading(f"Risk {risk_num}: {answer.question.content[:80]}", level=2)
            doc.add_paragraph(f"Kategori: {answer.question.topic.category.title}")
            doc.add_paragraph(f"Konu: {answer.question.topic.title}")
            if answer.risk_priority:
                doc.add_paragraph(f"Öncelik: {answer.get_risk_priority_display()}")
            
            if hasattr(answer, 'measures'):
                measures = answer.measures.all()
                if measures.exists():
                    doc.add_paragraph("Önlemler:")
                    for m in measures:
                        if m.description:
                            doc.add_paragraph(f"• {m.description}", style='List Bullet')
                        if m.responsible_person:
                            doc.add_paragraph(f"  Sorumlu: {m.responsible_person}")
                        if m.planning_end_date:
                            doc.add_paragraph(f"  Hedef Tarih: {m.planning_end_date.strftime('%d.%m.%Y')}")
            
            risk_num += 1
        
        # Custom risks
        for cr in session.custom_risks.filter(is_acceptable=False):
            doc.add_heading(f"Ek Risk Ω{risk_num}: {cr.description[:80]}", level=2)
            if cr.risk_priority:
                doc.add_paragraph(f"Öncelik: {cr.get_risk_priority_display()}")
            
            if hasattr(cr, 'custom_measures'):
                measures = cr.custom_measures.all()
                if measures.exists():
                    doc.add_paragraph("Önlemler:")
                    for m in measures:
                        if m.description:
                            doc.add_paragraph(f"• {m.description}", style='List Bullet')
            
            risk_num += 1
        
        # Signature Section
        doc.add_page_break()
        doc.add_heading('İmza Sirküleri', level=1)
        doc.add_paragraph("Bu risk değerlendirme raporu aşağıdaki kişiler tarafından onaylanmıştır.")
        doc.add_paragraph()
        
        # Add team members if they exist
        team_members = session.team_members.all()
        if team_members.exists():
            for member in team_members:
                p = doc.add_paragraph()
                p.add_run(f"\n{member.get_role_display()}: {member.name}").bold = True
                if member.title:
                    p.add_run(f" ({member.title})")
                p.add_run("\n\nİmza: _______________________      Tarih: _______________\n")
        else:
            # Default signature blocks
            for role in ['İşveren / Vekili', 'İSG Uzmanı', 'İş Yeri Hekimi', 'Çalışan Temsilcisi']:
                p = doc.add_paragraph()
                p.add_run(f"\n{role}:").bold = True
                p.add_run(" _______________________\n")
                p.add_run("\nİmza: _______________________      Tarih: _______________\n")
        
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="rapor_{session.pk}.docx"'
        doc.save(response)
        return response
        
    except Exception as e:
        return HttpResponse(f"Error generating Word document: {str(e)}", status=500)


@login_required
def export_report_pdf(request, session_pk):
    """Export risk overview as PDF using WeasyPrint"""
    import os
    from django.conf import settings
    
    try:
        session = get_object_or_404(AssessmentSession, pk=session_pk)
        
        # Font path for Turkish character support
        font_path = os.path.join(settings.BASE_DIR, 'static', 'fonts', 'Roboto-Regular.ttf')
        
        # Build rows
        rows_html = ""
        risk_num = 1
        
        for answer in session.answers.filter(response='NO').select_related('question'):
            priority_text = answer.get_risk_priority_display() if answer.risk_priority else "-"
            measures_count = answer.measures.count()
            rows_html += f"""
                <tr>
                    <td>R{risk_num}</td>
                    <td>{answer.question.content}</td>
                    <td>{priority_text}</td>
                    <td>{measures_count}</td>
                </tr>
            """
            risk_num += 1
        
        for cr in session.custom_risks.filter(is_acceptable=False):
            priority_text = cr.get_risk_priority_display() if cr.risk_priority else "-"
            measures_count = cr.custom_measures.count()
            rows_html += f"""
                <tr>
                    <td>Ω{risk_num}</td>
                    <td>{cr.description}</td>
                    <td>{priority_text}</td>
                    <td>{measures_count}</td>
                </tr>
            """
            risk_num += 1
        
        # Signature blocks
        signature_html = ""
        team_members = session.team_members.all()
        if team_members.exists():
            signature_html = "<h2>İmza Sirküleri</h2><table class='sig-table'>"
            for member in team_members:
                signature_html += f"""
                    <tr>
                        <td style="width:50%; padding-top:30px; border-bottom:1px solid #333;">
                            {member.name}<br><small>{member.get_role_display()}</small>
                        </td>
                        <td style="width:50%; padding-top:30px; border-bottom:1px solid #333;">
                            İmza: _______________________
                        </td>
                    </tr>
                """
            signature_html += "</table>"
        else:
            signature_html = """
            <h2>İmza Sirküleri</h2>
            <table class="sig-table">
                <tr><td>İşveren / Vekili</td><td>İmza: _______________________</td></tr>
                <tr><td>İSG Uzmanı</td><td>İmza: _______________________</td></tr>
                <tr><td>İş Yeri Hekimi</td><td>İmza: _______________________</td></tr>
                <tr><td>Çalışan Temsilcisi</td><td>İmza: _______________________</td></tr>
            </table>
            """
        
        # Participants section
        participants_html = ""
        if session.participants:
            participants_html = f"<p><strong>Katılımcılar:</strong> {session.participants}</p>"
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                @font-face {{
                    font-family: 'TurkishFont';
                    src: url('file://{font_path}');
                }}
                @page {{ size: A4; margin: 2cm; }}
                body {{ font-family: 'TurkishFont', sans-serif; font-size: 12px; }}
                h1 {{ color: #333; font-size: 22px; border-bottom: 2px solid #667eea; padding-bottom: 10px; }}
                h2 {{ color: #667eea; font-size: 16px; margin-top: 25px; }}
                table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
                th, td {{ border: 1px solid #ccc; padding: 8px; text-align: left; font-size: 10px; }}
                th {{ background-color: #667eea; color: white; }}
                .sig-table td {{ border: none; padding-top: 30px; }}
            </style>
        </head>
        <body>
            <h1>Risk Değerlendirme Özeti</h1>
            <p><strong>Tesis:</strong> {session.facility.name}</p>
            <p><strong>Araç:</strong> {session.tool.title}</p>
            <p><strong>Tarih:</strong> {session.created_at.strftime('%d.%m.%Y')}</p>
            {participants_html}
            
            <h2>Tespit Edilen Riskler</h2>
            <table>
                <tr><th>No</th><th>Risk Tanımı</th><th>Öncelik</th><th>Önlem Sayısı</th></tr>
                {rows_html}
            </table>
            
            {signature_html}
        </body>
        </html>
        """
        
        # Generate PDF with WeasyPrint
        try:
            from weasyprint import HTML
            pdf_file = HTML(string=html_content).write_pdf()
            response = HttpResponse(pdf_file, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="risk_ozet_{session.pk}.pdf"'
            return response
        except ImportError:
            return HttpResponse("Error: WeasyPrint is not installed. Please run: pip install weasyprint", status=500)
        except Exception as e:
            return HttpResponse(f"WeasyPrint error: {str(e)}", status=500)
                
    except Exception as e:
        return HttpResponse(f"Error generating PDF: {str(e)}", status=500)


@login_required
def export_full_checklist_pdf(request, session_pk):
    """Export full checklist with all questions and answers"""
    import os
    from django.conf import settings
    
    try:
        session = get_object_or_404(AssessmentSession, pk=session_pk)
        font_path = os.path.join(settings.BASE_DIR, 'static', 'fonts', 'DejaVuSans.ttf')
        
        # Get all questions for this tool
        questions = RiskQuestion.objects.filter(topic__category__tool=session.tool).select_related('topic__category')
        
        # Build rows
        rows_html = ""
        for idx, q in enumerate(questions, 1):
            answer = session.answers.filter(question=q).first()
            
            if answer and answer.response:
                if answer.response == 'YES':
                    status = '<span style="color: #28a745;">✓ Evet</span>'
                elif answer.response == 'NO':
                    status = '<span style="color: #dc3545;">✗ Hayır (Risk)</span>'
                elif answer.response == 'NA':
                    status = '<span style="color: #6c757d;">— İlgili Değil</span>'
                else:
                    status = answer.get_response_display()
                notes = answer.notes or ""
            else:
                status = '<span style="color: #adb5bd;">○ Cevaplanmadı</span>'
                notes = ""
            
            rows_html += f"""
                <tr>
                    <td>{idx}</td>
                    <td>{q.topic.category.title}</td>
                    <td>{q.content}</td>
                    <td>{status}</td>
                    <td>{notes}</td>
                </tr>
            """
        
        # Signature blocks
        signature_html = """
        <h2>İmza Sirküleri</h2>
        <table>
            <tr><td style="width:50%; padding-top:40px; border-bottom:1px solid #333;">İşveren / Vekili</td><td style="width:50%; padding-top:40px; border-bottom:1px solid #333;">İmza: _______________________</td></tr>
            <tr><td style="width:50%; padding-top:40px; border-bottom:1px solid #333;">İSG Uzmanı</td><td style="width:50%; padding-top:40px; border-bottom:1px solid #333;">İmza: _______________________</td></tr>
            <tr><td style="width:50%; padding-top:40px; border-bottom:1px solid #333;">İş Yeri Hekimi</td><td style="width:50%; padding-top:40px; border-bottom:1px solid #333;">İmza: _______________________</td></tr>
            <tr><td style="width:50%; padding-top:40px; border-bottom:1px solid #333;">Çalışan Temsilcisi</td><td style="width:50%; padding-top:40px; border-bottom:1px solid #333;">İmza: _______________________</td></tr>
        </table>
        """
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                @font-face {{
                    font-family: 'DejaVu';
                    src: url('file://{font_path}');
                }}
                @page {{ size: A4 landscape; margin: 1.5cm; }}
                body {{ font-family: 'DejaVu', sans-serif; font-size: 10px; }}
                h1 {{ color: #333; font-size: 20px; border-bottom: 2px solid #667eea; padding-bottom: 10px; }}
                h2 {{ color: #667eea; font-size: 14px; margin-top: 20px; }}
                table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
                th, td {{ border: 1px solid #ccc; padding: 6px; text-align: left; font-size: 9px; }}
                th {{ background-color: #667eea; color: white; }}
            </style>
        </head>
        <body>
            <h1>Denetim Listesi (Full Checklist)</h1>
            <p><strong>Tesis:</strong> {session.facility.name} | <strong>Araç:</strong> {session.tool.title} | <strong>Tarih:</strong> {session.created_at.strftime('%d.%m.%Y')}</p>
            
            <table>
                <tr><th>No</th><th>Kategori</th><th>Soru</th><th>Durum</th><th>Notlar</th></tr>
                {rows_html}
            </table>
            
            {signature_html}
        </body>
        </html>
        """
        
        from xhtml2pdf import pisa
        from io import BytesIO
        
        result = BytesIO()
        pdf = pisa.pisaDocument(BytesIO(html_content.encode("UTF-8")), result)
        
        if pdf.err:
            return HttpResponse(f"PDF generation error: {pdf.err}", status=500)
        
        response = HttpResponse(result.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="denetim_listesi_{session.pk}.pdf"'
        return response
        
    except Exception as e:
        return HttpResponse(f"Error generating PDF: {str(e)}", status=500)


@login_required
def assessment_team(request, session_pk):
    """Manage risk assessment team members"""
    session = get_object_or_404(AssessmentSession, pk=session_pk)
    
    if request.method == 'POST':
        action = request.POST.get('action')
        
        if action == 'add':
            role = request.POST.get('role')
            name = request.POST.get('name')
            title = request.POST.get('title', '')
            if role and name:
                RiskAssessmentTeamMember.objects.create(
                    session=session,
                    role=role,
                    name=name,
                    title=title
                )
                messages.success(request, 'Ekip üyesi eklendi.')
        
        elif action == 'delete':
            member_id = request.POST.get('member_id')
            if member_id:
                RiskAssessmentTeamMember.objects.filter(pk=member_id, session=session).delete()
                messages.success(request, 'Ekip üyesi silindi.')
        
        return redirect('assessment_team', session_pk=session.pk)
    
    team_members = session.team_members.all()
    
    context = {
        'session': session,
        'facility': session.facility,
        'tool': session.tool,
        'team_members': team_members,
        'role_choices': RiskAssessmentTeamMember.ROLE_CHOICES,
    }
    return render(request, 'core/assessment_team.html', context)


@login_required
def assessment_list(request):
    """List all risk assessments"""
    # Get filter parameters
    status_filter = request.GET.get('status', '')
    facility_filter = request.GET.get('facility', '')
    search = request.GET.get('search', '')
    
    assessments = AssessmentSession.objects.all().select_related('facility', 'tool', 'facility__workplace')
    
    if status_filter:
        assessments = assessments.filter(status=status_filter)
    
    if facility_filter:
        assessments = assessments.filter(facility_id=facility_filter)
    
    if search:
        assessments = assessments.filter(
            Q(title__icontains=search) | 
            Q(facility__name__icontains=search) |
            Q(tool__title__icontains=search)
        )
    
    facilities = Facility.objects.all()
    
    context = {
        'assessments': assessments,
        'facilities': facilities,
        'status_filter': status_filter,
        'facility_filter': facility_filter,
        'search': search,
    }
    return render(request, 'core/assessment_list.html', context)


@login_required
def assessment_delete(request, pk):
    """Delete a risk assessment"""
    assessment = get_object_or_404(AssessmentSession, pk=pk)
    
    if request.method == 'POST':
        title = assessment.title
        facility = assessment.facility
        assessment.delete()
        log_action(request.user, 'Silme', f"Risk Değerlendirmesi: {title}")
        messages.success(request, f'"{title}" değerlendirmesi silindi.')
        return redirect('assessment_list')
    
    return redirect('assessment_list')


@login_required
def assessment_bulk_delete(request):
    """Bulk delete assessments"""
    if request.method == 'POST':
        ids = request.POST.getlist('selected_ids')
        if ids:
            count = AssessmentSession.objects.filter(pk__in=ids).delete()[0]
            messages.success(request, f'{count} değerlendirme silindi.')
    return redirect('assessment_list')


@login_required
def education_certificate_word(request):
    """Export education certificate as Word document - matching PDF layout exactly"""
    education_id = request.GET.get('education_id')
    if not education_id:
        return HttpResponse("Error: education_id parameter required", status=400)
    
    from .models import Education, CertificateTemplate
    
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor, Twips
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.enum.section import WD_ORIENT
        from docx.oxml.ns import qn, nsdecls
        from docx.oxml import parse_xml
    except ImportError:
        return HttpResponse("Error: python-docx library is not installed. Please run: pip install python-docx", status=500)
    
    try:
        education = get_object_or_404(Education, pk=education_id)
        
        # Get template settings
        try:
            template_obj = CertificateTemplate.objects.get(name="Global")
            institute_lines = template_obj.institute_name.split('\\n')
            topics_list = [t.strip() for t in template_obj.education_topics.split('\n') if t.strip()]
        except CertificateTemplate.DoesNotExist:
            institute_lines = ["Kurum Adı Girilmedi"]
            topics_list = ["Konular Girilmedi"]
        
        # Get professionals
        specialist_name = ""
        medic_name = ""
        for p in education.professionals.all():
            if p.role == 'SPECIALIST':
                specialist_name = p.name
            elif p.role in ['DOCTOR', 'OTHER_HEALTH']:
                medic_name = p.name
        
        workers = education.workers.all()
        date_str = education.date.strftime('%d.%m.%Y')
        duration_str = f"{education.duration} Saat"
        
        doc = Document()
        
        # Set up page layout (A4, narrow margins)
        for section in doc.sections:
            section.page_width = Inches(8.27)  # A4 width
            section.page_height = Inches(11.69)  # A4 height
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(1.5)
        
        # Define colors
        RED = RGBColor(198, 40, 40)  # #C62828
        BLUE = RGBColor(31, 58, 88)  # #1F3A58
        GRAY = RGBColor(85, 85, 85)  # #555555
        
        for idx, worker in enumerate(workers):
            # === HEADER (Institution name) ===
            for line in institute_lines:
                header = doc.add_paragraph()
                header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                header_run = header.add_run(line.strip())
                header_run.font.size = Pt(10)
                header_run.font.bold = True
                header_run.font.color.rgb = RED
                header.paragraph_format.space_after = Pt(0)
                header.paragraph_format.space_before = Pt(0)
            
            # Add spacing after header
            doc.add_paragraph().paragraph_format.space_after = Pt(10)
            
            # === TITLE ===
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.add_run("EĞİTİM BELGESİ")
            title_run.font.size = Pt(28)
            title_run.font.bold = True
            title_run.font.color.rgb = BLUE
            title_run.font.underline = True
            title.paragraph_format.space_after = Pt(20)
            
            # === INFO SECTION ===
            info_data = [
                ("Sayı", f"{education.id}-{worker.id}"),
                ("TCKN", worker.tckn),
                ("Tarih", date_str),
                ("Süre", duration_str),
                ("İş Yeri", education.workplace.name.upper()),
            ]
            
            for label, value in info_data:
                info_p = doc.add_paragraph()
                label_run = info_p.add_run(f"{label}".ljust(10))
                label_run.font.size = Pt(11)
                label_run.font.bold = True
                colon_run = info_p.add_run(": ")
                colon_run.font.size = Pt(11)
                colon_run.font.bold = True
                value_run = info_p.add_run(str(value))
                value_run.font.size = Pt(11)
                value_run.font.bold = True
                info_p.paragraph_format.space_after = Pt(2)
                info_p.paragraph_format.space_before = Pt(0)
            
            # Add spacing
            doc.add_paragraph().paragraph_format.space_after = Pt(10)
            
            # === WORKER NAME (with underline) ===
            worker_p = doc.add_paragraph()
            worker_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            worker_run = worker_p.add_run(worker.name.upper())
            worker_run.font.size = Pt(18)
            worker_run.font.bold = True
            worker_run.font.color.rgb = BLUE
            worker_p.paragraph_format.space_after = Pt(15)
            
            # Add horizontal line under name
            line_p = doc.add_paragraph()
            line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            line_run = line_p.add_run("_" * 50)
            line_run.font.size = Pt(10)
            line_run.font.color.rgb = GRAY
            line_p.paragraph_format.space_after = Pt(15)
            
            # === BODY TEXT ===
            body = doc.add_paragraph()
            body.alignment = WD_ALIGN_PARAGRAPH.CENTER
            body.add_run("Yukarıda adı geçen çalışan,").font.size = Pt(10)
            body.paragraph_format.space_after = Pt(2)
            
            body2 = doc.add_paragraph()
            body2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            body2_run = body2.add_run("Çalışanların İş Sağlığı ve Güvenliği Eğitimlerinin Usul ve Esasları Hakkında Yönetmelik kapsamında verilen örgün ")
            body2_run.font.size = Pt(10)
            bold_run = body2.add_run("İş Sağlığı ve Güvenliği Eğitimini")
            bold_run.font.size = Pt(10)
            bold_run.font.bold = True
            body2.paragraph_format.space_after = Pt(2)
            
            body3 = doc.add_paragraph()
            body3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            body3_run = body3.add_run("başarıyla tamamlayarak bu belgeyi almaya hak kazanmıştır.")
            body3_run.font.size = Pt(10)
            body3.paragraph_format.space_after = Pt(20)
            
            # === TOPICS TITLE ===
            topics_title = doc.add_paragraph()
            topics_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            topics_run = topics_title.add_run("EĞİTİM KONULARI")
            topics_run.font.size = Pt(11)
            topics_run.font.bold = True
            topics_run.font.small_caps = True
            topics_title.paragraph_format.space_after = Pt(3)
            
            # === SEPARATOR ===
            sep_p = doc.add_paragraph()
            sep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sep_run = sep_p.add_run("~ ☚ ~")
            sep_run.font.size = Pt(14)
            sep_run.font.color.rgb = RGBColor(253, 216, 53)  # Gold
            sep_p.paragraph_format.space_after = Pt(10)
            
            # === TOPICS TABLE (2 columns) ===
            mid = (len(topics_list) + 1) // 2
            left_topics = topics_list[:mid]
            right_topics = topics_list[mid:]
            
            topics_table = doc.add_table(rows=max(len(left_topics), len(right_topics)), cols=2)
            topics_table.autofit = True
            
            for row_idx, row in enumerate(topics_table.rows):
                for col_idx, cell in enumerate(row.cells):
                    topics_to_use = left_topics if col_idx == 0 else right_topics
                    if row_idx < len(topics_to_use):
                        cell.text = topics_to_use[row_idx]
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.size = Pt(8)
                            para.paragraph_format.space_after = Pt(1)
            
            # Add spacing after topics
            doc.add_paragraph().paragraph_format.space_after = Pt(25)
            
            # === SIGNATURE TABLE ===
            sig_table = doc.add_table(rows=3, cols=3)
            sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            sig_data = [
                ["_______________________", "_______________________", "_______________________"],
                ["İş Güvenliği Uzmanı", "İş Yeri Hekimi/Hemşiresi", "İşveren/İşveren Vekili"],
                [specialist_name or "EĞİTMEN", medic_name or "", ""],
            ]
            
            for row_idx, row_data in enumerate(sig_data):
                for col_idx, cell_text in enumerate(row_data):
                    cell = sig_table.cell(row_idx, col_idx)
                    cell.text = cell_text
                    for para in cell.paragraphs:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.font.size = Pt(9)
                            if row_idx == 1:  # Role labels
                                run.font.bold = False
                                run.font.color.rgb = GRAY
                            elif row_idx == 2:  # Names
                                run.font.color.rgb = BLUE
            
            # Page break for next worker (except last)
            if idx < len(workers) - 1:
                doc.add_page_break()
        
        # Generate response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="egitim_belgesi_{education.id}.docx"'
        doc.save(response)
        return response
        
    except Exception as e:
        import traceback
        return HttpResponse(f"Error generating Word document: {str(e)}\n\n{traceback.format_exc()}", status=500)


# =============================================================================
# Fast Track Assessment Mode
# =============================================================================

from .risk_library import get_risk_library, get_risk_categories, search_risks

@login_required
def assessment_fast_run(request, pk):
    """Fast Track assessment runner - select risks from library"""
    session = get_object_or_404(AssessmentSession, pk=pk)
    
    # Get categories for filter dropdown
    categories = get_risk_categories()
    
    # Get existing custom risks for this session
    added_risks = session.custom_risks.all().prefetch_related('custom_measures', 'control_records')
    
    # Serialize added risks as JSON for safe JavaScript consumption
    import json
    
    def get_latest_residual(risk):
        """Get the residual score from the most recent control record"""
        latest = risk.control_records.order_by('-control_date', '-created_at').first()
        return latest.residual_score if latest else None
    
    added_risks_json = json.dumps([{
        'id': risk.pk,
        'libraryId': risk.source_library_id,
        'description': risk.description or '',
        'category': risk.category or '',
        'legal_basis': risk.legal_basis or '',
        'measure': risk.measure or '',
        'affected_persons': risk.affected_persons or '',
        'due_date': risk.due_date.isoformat() if risk.due_date else '',
        'responsible': risk.responsible_person or '',
        'mitigation_strategy': risk.mitigation_strategy or '',
        'estimated_budget': str(risk.estimated_budget) if risk.estimated_budget else '',
        'kinney_probability': risk.kinney_probability,
        'kinney_frequency': risk.kinney_frequency,
        'kinney_severity': risk.kinney_severity,
        'kinney_score': risk.kinney_score,
        'matrix_probability': risk.matrix_probability,
        'matrix_severity': risk.matrix_severity,
        'matrix_score': risk.matrix_score,
        'has_control_records': risk.control_records.exists(),
        'control_records_count': risk.control_records.count(),
        'latest_residual_score': get_latest_residual(risk),
    } for risk in added_risks])
    
    context = {
        'session': session,
        'facility': session.facility,
        'tool': session.tool,
        'title': session.title,
        'categories': categories,
        'added_risks': added_risks,
        'added_risks_json': added_risks_json,
        'added_risks_count': added_risks.count(),
    }
    return render(request, 'core/assessment_fast_run.html', context)


@login_required
def api_get_risk_library(request):
    """API endpoint to get paginated risk library"""
    query = request.GET.get('q', '')
    category = request.GET.get('category', '')
    limit = int(request.GET.get('limit', 50))
    offset = int(request.GET.get('offset', 0))
    
    result = search_risks(query=query, category=category, limit=limit, offset=offset)
    
    # Format results for JSON response
    formatted_results = []
    for risk in result['results']:
        formatted_results.append({
            'id': risk.get('id', 0),
            'grup_adi': risk.get('Grup Adı', ''),
            'tehlike': risk.get('Tehlike', ''),
            'risk': risk.get('Risk', ''),
            'mevzuat': risk.get('İlgili Mevzuat', ''),
            'onlem': risk.get('Alınması Gereken Önlemler', ''),
            'konu': risk.get('Konu', ''),
        })
    
    return JsonResponse({
        'results': formatted_results,
        'total': result['total'],
        'has_more': result['has_more'],
    })


@login_required
def api_get_risk_categories(request):
    """API endpoint to get risk categories"""
    categories = get_risk_categories()
    return JsonResponse({'categories': categories})


@login_required
def api_add_library_risk(request, session_pk):
    """API endpoint to add a risk from the library to the session"""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    
    session = get_object_or_404(AssessmentSession, pk=session_pk)
    
    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    
    library_id = data.get('library_id')
    
    # Get risk from library
    all_risks = get_risk_library()
    risk_data = None
    for risk in all_risks:
        if risk.get('id') == library_id:
            risk_data = risk
            break
    
    if not risk_data:
        return JsonResponse({'error': 'Risk not found in library'}, status=404)
    
    # Create custom risk with full category data
    tehlike = risk_data.get('Tehlike', '')
    risk_text = risk_data.get('Risk', '')
    description = f"{tehlike} - {risk_text}" if tehlike and risk_text else (tehlike or risk_text)
    
    # Get measure and affected persons from library
    onlem = risk_data.get('Alınması Gereken Önlemler', '')
    etkilenen = risk_data.get('Etkilenecek Kişiler', '')
    
    custom_risk = AssessmentCustomRisk.objects.create(
        session=session,
        description=description,
        is_acceptable=False,  # Risks need action
        category=risk_data.get('Grup Adı', '') or 'Genel Riskler',
        sub_category=risk_data.get('Üst Grup Adı', ''),
        hazard_source=risk_data.get('Tehlike Kaynağı', '') or tehlike,
        legal_basis=risk_data.get('İlgili Mevzuat', ''),
        source_library_id=library_id,
        affected_persons=etkilenen,
        measure=onlem,
    )
    
    # Also create action plan measure if available
    if onlem:
        ActionPlanMeasure.objects.create(
            custom_risk=custom_risk,
            description=onlem,
        )
    
    return JsonResponse({
        'success': True,
        'risk_id': custom_risk.pk,
        'description': custom_risk.description,
        'category': custom_risk.category,
        'sub_category': custom_risk.sub_category,
        'hazard_source': custom_risk.hazard_source,
        'legal_basis': custom_risk.legal_basis,
        'measure': onlem,
        'affected_persons': etkilenen,
    })


@login_required
def api_update_fast_risk(request, session_pk, risk_pk):
    """API endpoint to update a fast track risk"""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    
    session = get_object_or_404(AssessmentSession, pk=session_pk)
    custom_risk = get_object_or_404(AssessmentCustomRisk, pk=risk_pk, session=session)
    
    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    
    # Update risk fields
    if 'description' in data:
        custom_risk.description = data['description']
    if 'priority' in data:
        custom_risk.risk_priority = data['priority']
    if 'category' in data:
        custom_risk.category = data['category']
    if 'measure' in data:
        custom_risk.measure = data['measure']
    if 'responsible' in data:
        custom_risk.responsible_person = data['responsible']
    if 'affected_persons' in data:
        custom_risk.affected_persons = data['affected_persons']
    if 'due_date' in data:
        due_str = data['due_date']
        if due_str:
            try:
                from datetime import datetime
                custom_risk.due_date = datetime.strptime(due_str, '%Y-%m-%d').date()
            except ValueError:
                pass
        else:
            custom_risk.due_date = None
    
    # DÖF fields
    if 'mitigation_strategy' in data:
        custom_risk.mitigation_strategy = data['mitigation_strategy']
    if 'estimated_budget' in data:
        budget_str = data['estimated_budget']
        if budget_str:
            try:
                from decimal import Decimal
                custom_risk.estimated_budget = Decimal(budget_str.replace(',', '.'))
            except:
                pass
        else:
            custom_risk.estimated_budget = None
    
    # Scoring fields
    if 'scoring_method' in data:
        custom_risk.scoring_method = data['scoring_method']
    
    # Kinney scoring
    if 'kinney_probability' in data:
        custom_risk.kinney_probability = float(data['kinney_probability']) if data['kinney_probability'] else None
    if 'kinney_frequency' in data:
        custom_risk.kinney_frequency = float(data['kinney_frequency']) if data['kinney_frequency'] else None
    if 'kinney_severity' in data:
        custom_risk.kinney_severity = int(data['kinney_severity']) if data['kinney_severity'] else None
    
    # Matrix scoring
    if 'matrix_probability' in data:
        custom_risk.matrix_probability = int(data['matrix_probability']) if data['matrix_probability'] else None
    if 'matrix_severity' in data:
        custom_risk.matrix_severity = int(data['matrix_severity']) if data['matrix_severity'] else None
    
    custom_risk.save()  # Auto-calculates scores in save()
    
    # Also update or create ActionPlanMeasure for backwards compatibility
    if 'measure' in data or 'responsible' in data:
        measure = custom_risk.custom_measures.first()
        if measure:
            if 'measure' in data:
                measure.description = data['measure']
            if 'responsible' in data:
                measure.responsible_person = data['responsible']
            measure.save()
        elif data.get('measure') or data.get('responsible'):
            ActionPlanMeasure.objects.create(
                custom_risk=custom_risk,
                description=data.get('measure', ''),
                responsible_person=data.get('responsible', ''),
            )
    
    # Get risk level
    risk_level, risk_label = custom_risk.risk_level
    
    return JsonResponse({
        'success': True,
        'risk_id': custom_risk.pk,
        'kinney_score': custom_risk.kinney_score,
        'matrix_score': custom_risk.matrix_score,
        'risk_level': risk_level,
        'risk_label': risk_label,
    })


@login_required
def api_delete_fast_risk(request, session_pk, risk_pk):
    """API endpoint to delete a fast track risk"""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    
    session = get_object_or_404(AssessmentSession, pk=session_pk)
    custom_risk = get_object_or_404(AssessmentCustomRisk, pk=risk_pk, session=session)
    
    custom_risk.delete()
    
    return JsonResponse({
        'success': True,
    })


@login_required
def api_get_control_records(request, session_pk, risk_pk):
    """API endpoint to get all control records for a risk"""
    session = get_object_or_404(AssessmentSession, pk=session_pk)
    custom_risk = get_object_or_404(AssessmentCustomRisk, pk=risk_pk, session=session)
    
    records = []
    for record in custom_risk.control_records.all():
        risk_level, risk_label = record.risk_level
        records.append({
            'id': record.pk,
            'control_date': record.control_date.strftime('%Y-%m-%d'),
            'control_date_display': record.control_date.strftime('%d.%m.%Y'),
            'auditor_name': record.auditor_name,
            'observation_note': record.observation_note,
            'kinney_probability': record.kinney_probability,
            'kinney_frequency': record.kinney_frequency,
            'kinney_severity': record.kinney_severity,
            'residual_score': record.residual_score,
            'risk_level': risk_level,
            'risk_label': risk_label,
        })
    
    return JsonResponse({
        'success': True,
        'records': records,
        'original_score': custom_risk.kinney_score,
    })


@login_required
def api_create_control_record(request, session_pk, risk_pk):
    """API endpoint to create a new control/audit record for a risk"""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    
    session = get_object_or_404(AssessmentSession, pk=session_pk)
    custom_risk = get_object_or_404(AssessmentCustomRisk, pk=risk_pk, session=session)
    
    try:
        data = json.loads(request.body)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    
    # Parse control date
    control_date_str = data.get('control_date')
    if not control_date_str:
        return JsonResponse({'error': 'control_date is required'}, status=400)
    
    try:
        from datetime import datetime as dt
        control_date = dt.strptime(control_date_str, '%Y-%m-%d').date()
    except ValueError:
        return JsonResponse({'error': 'Invalid date format'}, status=400)
    
    auditor_name = data.get('auditor_name', '')
    if not auditor_name:
        return JsonResponse({'error': 'auditor_name is required'}, status=400)
    
    # Get scoring method (from request or from session)
    scoring_method = data.get('scoring_method', session.scoring_method)
    
    # Create the control record
    record = RiskControlRecord.objects.create(
        risk=custom_risk,
        control_date=control_date,
        auditor_name=auditor_name,
        observation_note=data.get('observation_note', ''),
        scoring_method=scoring_method,
        # Kinney fields
        kinney_probability=float(data['kinney_probability']) if data.get('kinney_probability') else None,
        kinney_frequency=float(data['kinney_frequency']) if data.get('kinney_frequency') else None,
        kinney_severity=int(data['kinney_severity']) if data.get('kinney_severity') else None,
        # Matrix fields
        matrix_probability=int(data['matrix_probability']) if data.get('matrix_probability') else None,
        matrix_severity=int(data['matrix_severity']) if data.get('matrix_severity') else None,
    )
    
    risk_level, risk_label = record.risk_level
    
    # Get appropriate original score based on scoring method
    original_score = custom_risk.matrix_score if scoring_method == 'MATRIX' else custom_risk.kinney_score
    
    return JsonResponse({
        'success': True,
        'record': {
            'id': record.pk,
            'control_date': record.control_date.strftime('%Y-%m-%d'),
            'control_date_display': record.control_date.strftime('%d.%m.%Y'),
            'auditor_name': record.auditor_name,
            'observation_note': record.observation_note,
            'kinney_probability': record.kinney_probability,
            'kinney_frequency': record.kinney_frequency,
            'kinney_severity': record.kinney_severity,
            'matrix_probability': record.matrix_probability,
            'matrix_severity': record.matrix_severity,
            'residual_score': record.residual_score,
            'risk_level': risk_level,
            'risk_label': risk_label,
        },
        'original_score': original_score,
    })


# =============================================================================
# Public Safety Forum (QR Code Access)
# =============================================================================

from .forms import PublicEngagementForm
from .models import SafetyEngagement, SafetyPoll


def public_safety_forum(request, facility_uuid):
    """Public forum view accessible via QR code - NO LOGIN REQUIRED"""
    facility = get_object_or_404(Facility, uuid=facility_uuid)
    
    # Get form for submissions
    form = PublicEngagementForm(request=request)
    
    # Get published engagements for the community wall
    wall_items = SafetyEngagement.objects.filter(
        facility=facility,
        is_public_on_wall=True
    ).order_by('-created_at')[:20]
    
    # Get active polls
    polls = SafetyPoll.objects.filter(
        facility=facility,
        is_active=True
    ).order_by('-created_at')
    
    # Check which polls user has already voted on (cookie-based)
    voted_polls = request.COOKIES.get('voted_polls', '').split(',')
    
    # Get captcha for comment forms
    captcha_question = request.session.get('safety_math_question', '3 + 5 = ?')
    captcha_answer = request.session.get('safety_math_answer', 8)
    
    context = {
        'facility': facility,
        'form': form,
        'wall_items': wall_items,
        'polls': polls,
        'voted_polls': voted_polls,
        'success_message': request.session.pop('forum_success', None),
        'captcha_question': captcha_question,
        'captcha_answer': captcha_answer,
    }
    return render(request, 'core/public_forum.html', context)


def public_safety_submit(request, facility_uuid):
    """Handle public form submission with honeypot + math captcha"""
    if request.method != 'POST':
        return redirect('public_safety_forum', facility_uuid=facility_uuid)
    
    facility = get_object_or_404(Facility, uuid=facility_uuid)
    form = PublicEngagementForm(request.POST, request=request)
    
    if form.is_valid():
        # Create the engagement with PENDING status (requires approval)
        SafetyEngagement.objects.create(
            facility=facility,
            worker=None,  # Anonymous
            topic=form.cleaned_data['topic'],
            message=form.cleaned_data['message'],
            is_anonymous=True,
            status='PENDING',  # Requires professional approval
        )
        
        # Clear captcha for next submission
        request.session.pop('safety_math_answer', None)
        request.session.pop('safety_math_question', None)
        
        # Set success message - let user know it's pending
        request.session['forum_success'] = 'Bildiriminiz alındı ve onay bekliyor. Teşekkür ederiz!'
        
        return redirect('public_safety_forum', facility_uuid=facility_uuid)
    
    # If form invalid, re-render with errors
    wall_items = SafetyEngagement.objects.filter(
        facility=facility,
        is_public_on_wall=True
    ).order_by('-created_at')[:20]
    
    polls = SafetyPoll.objects.filter(
        facility=facility,
        is_active=True
    ).order_by('-created_at')
    
    voted_polls = request.COOKIES.get('voted_polls', '').split(',')
    
    # Get captcha for re-render
    captcha_question = request.session.get('safety_math_question', '3 + 5 = ?')
    captcha_answer = request.session.get('safety_math_answer', 8)
    
    context = {
        'facility': facility,
        'form': form,
        'wall_items': wall_items,
        'polls': polls,
        'voted_polls': voted_polls,
        'captcha_question': captcha_question,
        'captcha_answer': captcha_answer,
    }
    return render(request, 'core/public_forum.html', context)


def public_poll_vote(request, facility_uuid, poll_id):
    """Handle poll voting with cookie-based duplicate prevention"""
    if request.method != 'POST':
        return redirect('public_safety_forum', facility_uuid=facility_uuid)
    
    facility = get_object_or_404(Facility, uuid=facility_uuid)
    poll = get_object_or_404(SafetyPoll, pk=poll_id, facility=facility, is_active=True)
    
    # Check if already voted (cookie-based)
    voted_polls = request.COOKIES.get('voted_polls', '').split(',')
    if str(poll_id) in voted_polls:
        return redirect('public_safety_forum', facility_uuid=facility_uuid)
    
    # Get the selected option
    option = request.POST.get('option')
    if option and option in poll.options:
        # Update votes
        votes = poll.votes or {}
        votes[option] = votes.get(option, 0) + 1
        poll.votes = votes
        poll.save()
    
    # Set cookie to prevent re-voting
    response = redirect('public_safety_forum', facility_uuid=facility_uuid)
    voted_polls.append(str(poll_id))
    response.set_cookie('voted_polls', ','.join(filter(None, voted_polls)), max_age=365*24*60*60)
    
    return response


@login_required
def facility_qr_code(request, pk):
    """Generate and serve QR code for facility's public forum URL"""
    facility = get_object_or_404(Facility, pk=pk)
    
    try:
        import qrcode
        from io import BytesIO
    except ImportError:
        return HttpResponse("QR code library not installed. Run: pip install qrcode[pil]", status=500)
    
    # Get the participation URL
    url = facility.get_participation_url()
    
    # Generate QR code
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=4,
    )
    qr.add_data(url)
    qr.make(fit=True)
    
    img = qr.make_image(fill_color="black", back_color="white")
    
    # Save to bytes
    buffer = BytesIO()
    img.save(buffer, format='PNG')
    buffer.seek(0)
    
    # Return as downloadable PNG
    response = HttpResponse(buffer.getvalue(), content_type='image/png')
    response['Content-Disposition'] = f'attachment; filename="qr_forum_{facility.pk}.png"'
    return response


@login_required
def facility_engagements(request, pk):
    """Management view for facility engagements"""
    facility = get_object_or_404(Facility, pk=pk)
    
    engagements = SafetyEngagement.objects.filter(facility=facility).order_by('-created_at')
    
    context = {
        'facility': facility,
        'workplace': facility.workplace,
        'engagements': engagements,
    }
    return render(request, 'core/facility_engagements.html', context)


@login_required
def toggle_engagement_wall(request, pk):
    """API to toggle is_public_on_wall"""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    
    engagement = get_object_or_404(SafetyEngagement, pk=pk)
    
    try:
        data = json.loads(request.body)
        engagement.is_public_on_wall = data.get('is_public', False)
        engagement.save()
        return JsonResponse({'success': True, 'is_public_on_wall': engagement.is_public_on_wall})
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)


@login_required
def respond_to_engagement(request, pk):
    """API to add a comment to an engagement (creates EngagementComment)"""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    
    engagement = get_object_or_404(SafetyEngagement, pk=pk)
    
    try:
        data = json.loads(request.body)
        text = data.get('response', '').strip()
        
        if text:
            from .models import EngagementComment
            
            # Determine if user is a professional
            is_professional = request.user.is_staff  # Staff users are professionals
            author_name = f"{request.user.first_name} {request.user.last_name}".strip() or request.user.username
            if is_professional:
                author_name += " (İGU)"
            
            EngagementComment.objects.create(
                engagement=engagement,
                author_name=author_name,
                text=text,
                is_professional=is_professional
            )
        
        if data.get('resolved'):
            from django.utils import timezone
            engagement.resolved_at = timezone.now()
            engagement.save()
        
        return JsonResponse({'success': True})
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)


from .forms import SafetyPollForm, EngagementCommentForm
from .models import EngagementComment


@login_required
def poll_create(request, facility_pk):
    """Create a new poll for a facility"""
    facility = get_object_or_404(Facility, pk=facility_pk)
    
    if request.method == 'POST':
        form = SafetyPollForm(request.POST)
        if form.is_valid():
            SafetyPoll.objects.create(
                facility=facility,
                question=form.cleaned_data['question'],
                options=form.cleaned_data['options'],
                votes={opt: 0 for opt in form.cleaned_data['options']}
            )
            messages.success(request, 'Anket başarıyla oluşturuldu.')
            return redirect('facility_engagements', pk=facility.pk)
    else:
        form = SafetyPollForm()
    
    context = {
        'facility': facility,
        'form': form,
    }
    return render(request, 'core/poll_create.html', context)


def api_poll_vote(request, poll_id):
    """AJAX endpoint for voting on a poll - requires captcha"""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    
    poll = get_object_or_404(SafetyPoll, pk=poll_id, is_active=True)
    
    try:
        data = json.loads(request.body)
        option = data.get('option')
        captcha_input = data.get('captcha')
        
        # Validate captcha
        expected_answer = request.session.get('safety_math_answer', 8)
        if captcha_input != expected_answer:
            return JsonResponse({'error': 'Güvenlik sorusu yanlış'}, status=400)
        
        if option and option in poll.options:
            votes = poll.votes or {}
            votes[option] = votes.get(option, 0) + 1
            poll.votes = votes
            poll.save()
            
            return JsonResponse({
                'success': True,
                'total_votes': poll.get_total_votes(),
                'percentages': poll.get_percentages(),
                'votes': poll.votes
            })
        
        return JsonResponse({'error': 'Invalid option'}, status=400)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)


def api_poll_results(request, poll_id):
    """Get poll results without voting - no captcha required"""
    poll = get_object_or_404(SafetyPoll, pk=poll_id, is_active=True)
    
    return JsonResponse({
        'success': True,
        'question': poll.question,
        'total_votes': poll.get_total_votes(),
        'percentages': poll.get_percentages(),
        'votes': poll.votes or {}
    })



@login_required
def add_engagement_comment(request, engagement_pk):
    """Add a comment to an engagement"""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    
    engagement = get_object_or_404(SafetyEngagement, pk=engagement_pk)
    
    try:
        data = json.loads(request.body)
        text = data.get('text', '').strip()
        
        if text:
            is_professional = request.user.is_staff
            author_name = f"{request.user.first_name} {request.user.last_name}".strip() or request.user.username
            if is_professional:
                author_name += " (İGU)"
            
            comment = EngagementComment.objects.create(
                engagement=engagement,
                author_name=author_name,
                text=text,
                is_professional=is_professional
            )
            
            return JsonResponse({
                'success': True,
                'comment': {
                    'id': comment.pk,
                    'author_name': comment.author_name,
                    'text': comment.text,
                    'is_professional': comment.is_professional,
                    'created_at': comment.created_at.strftime('%d.%m.%Y %H:%M')
                }
            })
        
        return JsonResponse({'error': 'Text required'}, status=400)
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)


def api_engagement_like(request, engagement_pk):
    """AJAX endpoint to upvote an engagement - one vote per user using cookies"""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    
    engagement = get_object_or_404(SafetyEngagement, pk=engagement_pk)
    
    # Check if user already upvoted using cookies
    upvoted_items = request.COOKIES.get('upvoted_engagements', '').split(',')
    engagement_key = str(engagement_pk)
    
    if engagement_key in upvoted_items:
        return JsonResponse({
            'success': False,
            'already_voted': True,
            'likes': engagement.likes,
            'message': 'Zaten oy verdiniz'
        })
    
    # Add upvote
    engagement.likes += 1
    engagement.save()
    
    # Update cookie
    upvoted_items.append(engagement_key)
    
    response = JsonResponse({
        'success': True,
        'likes': engagement.likes
    })
    response.set_cookie('upvoted_engagements', ','.join(filter(None, upvoted_items)), max_age=365*24*60*60)
    
    return response


def add_public_comment(request, engagement_pk):
    """Public API to add anonymous comment with honeypot + captcha"""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    
    engagement = get_object_or_404(SafetyEngagement, pk=engagement_pk)
    
    try:
        data = json.loads(request.body)
        
        # Honeypot check
        if data.get('website_url'):
            return JsonResponse({'error': 'Spam detected'}, status=400)
        
        # Captcha validation
        captcha_input = data.get('captcha')
        expected_answer = request.session.get('safety_math_answer', 8)
        
        if captcha_input != expected_answer:
            return JsonResponse({'error': 'Güvenlik sorusu yanlış'}, status=400)
        
        text = data.get('text', '').strip()
        
        if not text:
            return JsonResponse({'error': 'Yorum boş olamaz'}, status=400)
        
        # Create anonymous comment
        comment = EngagementComment.objects.create(
            engagement=engagement,
            author_name="Bir Çalışan",
            text=text,
            is_professional=False
        )
        
        return JsonResponse({
            'success': True,
            'comment': {
                'id': comment.pk,
                'author_name': comment.author_name,
                'text': comment.text,
                'created_at': comment.created_at.strftime('%d.%m.%Y %H:%M')
            }
        })
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)


def api_wall_items(request, facility_uuid):
    """JSON API endpoint for wall items - for client-side rendering"""
    facility = get_object_or_404(Facility, uuid=facility_uuid)
    
    # Include both APPROVED and PENDING items (PENDING shown with hidden content)
    from django.db.models import Q
    wall_items = SafetyEngagement.objects.filter(
        Q(facility=facility) & (
            Q(is_public_on_wall=True, status='APPROVED') |
            Q(status='PENDING')
        )
    ).order_by('-created_at')[:20]
    
    # Convert to JSON-serializable format
    items_data = []
    for item in wall_items:
        topic_labels = {
            'SUGGESTION': '💡 Öneri',
            'NEAR_MISS': '⚠️ Ramak Kala',
            'HAZARD': '🚨 Tehlike',
            'COMPLAINT': '💬 Şikayet'
        }
        
        is_pending = item.status == 'PENDING'
        
        # Get all comments that are marked as public on voice
        public_comments = []
        expert_comment = None
        if not is_pending:
            voice_comments = item.comments.filter(is_public_on_voice=True).order_by('-created_at')
            for comment in voice_comments:
                comment_data = {
                    'author': comment.author_name,
                    'text': comment.text[:150] + ('...' if len(comment.text) > 150 else ''),
                    'is_professional': comment.is_professional
                }
                public_comments.append(comment_data)
                # Set expert_comment for backwards compatibility (first professional comment)
                if comment.is_professional and not expert_comment:
                    expert_comment = {
                        'author': comment.author_name,
                        'text': comment.text[:100] + ('...' if len(comment.text) > 100 else '')
                    }
        
        items_data.append({
            'id': item.pk,
            'topic': item.topic,
            'topic_label': topic_labels.get(item.topic, '💬 Diğer') if not is_pending else '',
            'message': item.message[:200] + ('...' if len(item.message) > 200 else '') if not is_pending else '',
            'created_at': item.created_at.strftime('%d.%m.%Y'),
            'comments_count': item.comments.count() if not is_pending else 0,
            'likes': item.likes if not is_pending else 0,
            'has_expert_comment': bool(expert_comment) if not is_pending else False,
            'expert_comment': expert_comment,
            'public_comments': public_comments,  # All public comments
            'status': item.status,
            'is_pending': is_pending
        })
    
    return JsonResponse({'items': items_data})


def api_public_polls(request, facility_uuid):
    """JSON API endpoint for polls - for client-side rendering"""
    facility = get_object_or_404(Facility, uuid=facility_uuid)
    
    polls = SafetyPoll.objects.filter(
        facility=facility,
        is_active=True
    ).order_by('-created_at')
    
    # Check which polls user has already voted on (cookie-based)
    voted_polls = request.COOKIES.get('voted_polls', '').split(',')
    
    polls_data = []
    for poll in polls:
        polls_data.append({
            'id': poll.pk,
            'question': poll.question,
            'options': poll.options,
            'total_votes': poll.get_total_votes(),
            'percentages': poll.get_percentages(),
            'votes': poll.votes or {},
            'has_voted': str(poll.pk) in voted_polls
        })
    
    return JsonResponse({'polls': polls_data})


@login_required
def api_update_engagement_status(request, engagement_pk):
    """API to approve or reject an engagement - for professional users only"""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    
    engagement = get_object_or_404(SafetyEngagement, pk=engagement_pk)
    
    try:
        data = json.loads(request.body)
        new_status = data.get('status')
        
        if new_status not in ['APPROVED', 'REJECTED', 'PENDING']:
            return JsonResponse({'error': 'Invalid status'}, status=400)
        
        old_status = engagement.status
        engagement.status = new_status
        
        # If approving, also set is_public_on_wall
        if new_status == 'APPROVED':
            engagement.is_public_on_wall = True
        
        engagement.save()
        
        return JsonResponse({
            'success': True,
            'old_status': old_status,
            'new_status': new_status,
            'is_public_on_wall': engagement.is_public_on_wall
        })
    except json.JSONDecodeError:
        return JsonResponse({'error': 'Invalid JSON'}, status=400)


@login_required
def api_delete_engagement(request, engagement_pk):
    """API to delete an engagement - for professional users only"""
    if request.method != 'DELETE' and request.method != 'POST':
        return JsonResponse({'error': 'DELETE or POST required'}, status=405)
    
    engagement = get_object_or_404(SafetyEngagement, pk=engagement_pk)
    engagement_id = engagement.pk
    engagement.delete()
    
    return JsonResponse({
        'success': True,
        'deleted_id': engagement_id
    })


@login_required
def api_toggle_comment_voice(request, comment_pk):
    """API to toggle comment visibility on Voice public page"""
    if request.method != 'POST':
        return JsonResponse({'error': 'POST required'}, status=405)
    
    comment = get_object_or_404(EngagementComment, pk=comment_pk)
    comment.is_public_on_voice = not comment.is_public_on_voice
    comment.save()
    
    return JsonResponse({
        'success': True,
        'is_public_on_voice': comment.is_public_on_voice
    })


# Footer Pages
@login_required
def privacy_page(request):
    """Privacy policy page"""
    return render(request, 'core/privacy.html')


@login_required
def terms_page(request):
    """Terms of service page"""
    return render(request, 'core/terms.html')


@login_required
def support_page(request):
    """Technical support page"""
    return render(request, 'core/support.html')


# Dashboard Search API
@login_required
def api_dashboard_search(request):
    """Search workplaces and facilities for dashboard search bar"""
    query = request.GET.get('q', '').strip()
    
    if len(query) < 2:
        return JsonResponse({'results': []})
    
    results = []
    
    # Search Workplaces
    workplaces = Workplace.objects.filter(
        Q(name__icontains=query) | Q(detsis_number__icontains=query)
    )[:5]
    
    for wp in workplaces:
        results.append({
            'type': 'workplace',
            'id': wp.pk,
            'name': wp.name,
            'description': f"İş Yeri • {wp.get_hazard_class_display() if wp.hazard_class else 'Tehlike sınıfı belirtilmemiş'}",
            'url': f"/workplaces/{wp.pk}/"
        })
    
    # Search Facilities
    facilities = Facility.objects.filter(
        Q(name__icontains=query)
    ).select_related('workplace')[:5]
    
    for fac in facilities:
        results.append({
            'type': 'facility',
            'id': fac.pk,
            'name': fac.name,
            'description': f"Birim • {fac.workplace.name if fac.workplace else 'Bağlı iş yeri yok'}",
            'url': f"/facilities/{fac.pk}/"
        })
    
    return JsonResponse({'results': results})
